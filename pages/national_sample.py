import streamlit as st
import pandas as pd
import base64
import pydeck as pdk
import pandas as pd
import numpy as np
import streamlit as st
import base64
from docx import Document
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from urllib.parse import urlencode
from google.oauth2.service_account import Credentials
import gspread

st.markdown("""
    <div style='width: 100%; padding: 20px 30px; background: #ffffff;
                border-bottom: 1px solid #e6e6e6; display: flex;
                justify-content: space-between; align-items: center;'>
        <a href="/" style='font-size: 18px; font-weight: 600; color: #344b77;
                text-decoration: none;'>← Faqja kryesore</a>
    </div>
"""
            , unsafe_allow_html=True)

st.markdown("""
<style>
            
/* Buttons */
.stButton > button {
    padding: 10px 18px !important;
    font-size: 16px !important;
    font-weight: 600 !important;
    border-radius: 10px !important;
    border: 2px solid #344b77 !important;
    color: #344b77 !important;
    background: #fff !important;
    transition: 0.25s ease !important;
}
            
.stButton > button:hover {
    background: #344b77 !important;
    color: white !important;
}

/* Hide only page navigation links, keep widgets **/
[data-testid="stSidebarNav"] li {
    display: none !important;
}

/* Keep sidebar header + widgets */
[data-testid="stSidebarNav"] {
    padding-bottom: 0 !important;
}    

.card {
    width: 100%;
    min-height: 185px;
    padding: 15px 20px;
    border-radius: 12px;
    background-color: #ffffff;
    border: 1px solid #e6e6e6;
    box-shadow: 0 1px 3px rgba(0,0,0,0.07);
    margin-bottom: 10px;
    }
.card-title {
    font-size: 18px;
    font-weight: 600;
    color: #344b77;
    margin-bottom: 8px;
    display: flex;
    align-items: center;
    gap: 8px;
    }
.card-title svg {
    width: 20px;
    height: 20px;
    }
.card-value {
    font-size: 16px;
    color: #000000;
    margin-bottom: 4px;
    }  

/* Set sidebar width */
[data-testid="stSidebar"] {
    width: 25% !important;
    min-width: 25% !important;
}       
            
</style>
""", unsafe_allow_html=True)


narrative_template_common = """
**Methodology and Sampling Frame**

The National Social Survey for Kosovo is based on the principles of {survey_label} survey methodology. In more specific terms, the collection of data for this quantitative study is executed through {methodology_label}. The proposed sample of **N={n_total}** has a 95% of confidence level with a **{moe}** of margin of error.
The sample design is based on a modified multistage random sampling methodology. Crucially, the sampling frame is constructed using a dual-source approach to address recent demographic data challenges:

- **General Population**: For the majority of municipalities, the sample design is based on the preliminary results of the Kosovo Census 2024, conducted by the Kosovo Agency of Statistics (KAS).
- **Serb-Majority Municipalities**: Due to the significant non-participation of the ethnic Serb community in the 2024 Census, population estimates for Serb-majority municipalities are derived from the 2018 OSCE Municipal Profiles, which provide the most reliable current estimates for these specific areas.

**Stratification and Quotas**

Sample quotas are calculated maintaining Probability Proportionate to Size (PPS). The sampling design follows a hierarchical stratification logic to ensure representativeness. The levels of stratification for this specific survey are applied in the following order:

- **Primary Stratification**: {primary_level}
- **Secondary Stratification**: {second_level}
- **Tertiary Stratification**: {third_level}
"""

narrative_template_capi = """
For each municipality, the list of settlements from the combined KAS/OSCE data is used to randomly select the Primary Sampling Units (PSUs). The nominal number of interviews for a single PSU is set at **{interviews_per_psu} interviews**.
If a settlement is allocated more interviews than the nominal PSU size, additional PSUs are selected within that settlement. In larger urban areas (which KAS often classifies as a single settlement), a neighborhood-based stratification is applied. Selection of additional PSUs is organized by dividing the settlement according to the multiples of {interviews_per_psu}, following a counter-clockwise orientation from a central landmark.
"""

narrative_template_cati = """
For the telephone survey, the selection of the settlement/municipality is done via Random Digit Dialing (RDD) or database selection to comply with the probability sampling condition. The "PSU" in this context refers to the individual valid phone number, stratified by region to match the geographic quotas defined above.  
"""

narrative_template_oversampling = """
**Oversampling Procedures** 

To ensure reliable estimates for specific sub-groups that would otherwise have a remote chance of being interviewed at statistically significant levels, this survey employs oversampling.  
"""

narrative_template_oversampling_inactive = """
 For this specific project, no oversampling was applied; the sample strictly follows the proportional population distribution.  
"""

narrative_template_oversampling_single_active = """
For this project, we have oversampled the following categories:

- **Target Group:** {os_target_group}
- **Method:** We selected **{os_added_total}** additional respondents from this group beyond their natural population proportion.

Adjustments for this oversampling will be required at the analytical phase.  
Weights will be applied so that oversampled groups are brought back to their true share in the population (KAS 2024, OSCE 2024).  
This ensures that national-level estimates remain representative while still allowing reliable subgroup analysis.
"""

narrative_template_oversampling_multi_active = """
For this project, we have oversampled we have oversampled the following categories:

{os_group_list}

For these groups, the additional respondents were selected beyond their natural population proportions. The added interviews ensure sufficient analytical precision for each oversampled subgroup.

Adjustments for this oversampling will be required at the analytical phase.  
Weights will be applied so that oversampled groups are brought back to their true share in the population (KAS 2024, OSCE 2018).  
This ensures that national-level estimates remain representative while still allowing reliable subgroup analysis.
"""

RECODE_D3_TEMPLATE = r"""
RECODE D3 (2=7)
(3=4)
(8=1)
(4=5)
(5=1)
(6=7)
(7=6)
(9=1)
(10=5)
(11=3)
(12=7)
(13=5)
(14=6)
(15=3)
(16=6)
(17=2)
(18=1)
(19=4)
(20=4)
(21=2)
(22=2)
(23=1)
(24=1)
(25=6)
(26=3)
(27=1)
(1=1)
(28=4)
(29=7)
(30=6)
(33=5)
(34=5)
(31=2)
(32=4)
(35=6)
(36=2)
(37=2)
(38=2)
 INTO Regjioni.
VARIABLE LABELS  Regjioni 'Regjioni'.
EXECUTE.
"""

TRANSLATIONS = {

    # ====================
    # VARIABLES (Dimensions)
    # ====================
    "Komuna": "Municipality",
    "Komunë": "Municipality",
    "Regjion": "Region",
    "Vendbanimi": "Settlement",
    "Vendbanim" :"Settlement",
    "Etnia": "Ethnicity",
    "Etni": "Ethnicity",
    "Gjinia": "Gender",
    "Mosha": "Age group",
    "Grupmosha": "Age group",

    # ====================
    # VALUES: Gender
    # ====================
    "Femra": "Female",
    "Meshkuj": "Male",

    # ====================
    # VALUES: Ethnicity
    # ====================
    "Shqiptar": "Albanian",
    "Serb": "Serb",
    "Tjerë": "Other",

    # ====================
    # VALUES: Regions
    # ====================
    "Prishtinë": "Pristina",
    "Mitrovicë": "Mitrovica",
    "Gjilan": "Gjilan",
    "Gjakovë": "Gjakova",
    "Ferizaj": "Ferizaj",
    "Prizren": "Prizren",
    "Pejë": "Peja",
}



# =========================
# CONFIG
# =========================

st.set_page_config(
    page_title="Dizajnimi i Mostrës Nacionale",
    layout="wide"
)

# =========================
# HELPERS
# =========================

@st.cache_data
def load_ethnicity_settlement_data(path: str) -> pd.DataFrame:
    """
    Load ASK-2024-Komuna-Etnia-Vendbanimi.xlsx
    Expected structure:
    Komuna | Vendbanimi | Shqiptar | Tjerë | Serb | ...
    Convert to long format: one row per (Komuna, Vendbanimi, Etnia).
    """
    df = pd.read_excel(path, sheet_name=0)
    # Identify ethnicity columns (all non-id cols except Komuna, Vendbanimi)
    id_cols = ["Komuna", "Vendbanimi"]
    eth_cols = [c for c in df.columns if c not in id_cols]
    df_long = df.melt(
        id_vars=id_cols,
        value_vars=eth_cols,
        var_name="Etnia",
        value_name="Pop_base"
    )
    # Clean
    df_long["Etnia"] = df_long["Etnia"].str.strip()
    df_long["Vendbanimi"] = df_long["Vendbanimi"].str.strip()
    df_long["Komuna"] = df_long["Komuna"].str.strip()
    df_long["Pop_base"] = df_long["Pop_base"].fillna(0).astype(float)
    return df_long


@st.cache_data
def load_gender_age_data(path: str) -> pd.DataFrame:
    """
    Load ASK-2024-Komuna-Gjinia-Mosha.xlsx (sheet 'census_00').
    Expected structure:
    Komuna | Gjinia | 0 | 1 | 2 | ... | (age columns)
    """
    df = pd.read_excel(path, sheet_name="census_00")
    # Normalize names
    df["Komuna"] = df["Komuna"].astype(str).str.strip()
    df["Gjinia"] = df["Gjinia"].astype(str).str.strip()

    # Keep only age columns that are pure ints as strings, e.g. "0","1",...
    age_cols = []
    for c in df.columns:
        s = str(c).strip()
        if s.isdigit():
            age_cols.append(c)

    # Remove empty rows (if any)
    df = df.dropna(subset=["Komuna", "Gjinia"])
    return df, age_cols

@st.cache_data
def load_contacts(path: str) -> pd.DataFrame:
    df = pd.read_excel(path)

    # Standardize
    for c in df.columns:
        df[c] = df[c].astype(str).str.strip()

    return df

def get_region_mapping() -> dict:
    """
    Map Komuna -> Regjion (bazuar në ASK)
    """
    region_map = {
        "Deçan": "Gjakovë",
        "Dragash": "Prizren",
        "Ferizaj": "Ferizaj",
        "Fushë Kosovë": "Prishtinë",
        "Gjakovë": "Gjakovë",
        "Gjilan": "Gjilan",
        "Gllogoc": "Prishtinë",
        "Graçanicë": "Prishtinë",
        "Han i Elezit": "Ferizaj",
        "Istog": "Pejë",
        "Junik": "Gjakovë",
        "Kaçanik": "Ferizaj",
        "Kamenicë": "Gjilan",
        "Klinë": "Pejë",
        "Kllokot": "Gjilan",
        "Leposavic": "Mitrovicë",
        "Lipjan": "Prishtinë",
        "Malishevë": "Prizren",
        "Mamushë": "Prizren",
        "Mitrovicë": "Mitrovicë",
        "Mitrovica Veriore": "Mitrovicë",
        "Novobërdë": "Gjilan",
        "Obiliq": "Prishtinë",
        "Partesh": "Gjilan",
        "Pejë": "Pejë",
        "Podujevë": "Prishtinë",
        "Prishtinë": "Prishtinë",
        "Prizren": "Prizren",
        "Rahovec": "Gjakovë",
        "Ranillug": "Gjilan",
        "Shtërpcë": "Ferizaj",
        "Shtime": "Ferizaj",
        "Skënderaj": "Mitrovicë",
        "Suharekë": "Prizren",
        "Viti": "Gjilan",
        "Vushtrri": "Mitrovicë",
        "Zubin Potok": "Mitrovicë",
        "Zvecan": "Mitrovicë"
    }
    return region_map

def compute_gender_age_coefficients(df_ga: pd.DataFrame,
                                    age_cols,
                                    selected_genders,
                                    min_age: int,
                                    max_age: int | None) -> pd.Series:
    """
    Compute coefficient per Komuna:
    coef(komuna) = Pop_selected(komuna) / Pop_total(komuna)

    - Pop_selected: individuals that match selected genders & [min_age, max_age]
    - Pop_total: all genders, all ages
    """
    df = df_ga.copy()

    # Sort age columns numerically
    age_cols_sorted = sorted(age_cols, key=lambda x: int(str(x)))
    max_available_age = int(str(age_cols_sorted[-1]))

    if max_age is None:
        max_age = max_available_age

    min_age = int(min_age)
    max_age = int(max_age)

    # Selected ages
    selected_age_cols = [
        c for c in age_cols_sorted
        if min_age <= int(c) <= max_age
    ]

    if not selected_age_cols:
        # No matching age columns: coefficient 0 for all
        return pd.Series(
            0.0,
            index=df["Komuna"].unique()
        )

    # Total population (all genders, all ages)
    df["Pop_all_ages"] = df[age_cols_sorted].sum(axis=1)

    # Filtered population by age range
    df["Pop_age_range"] = df[selected_age_cols].sum(axis=1)

    def calc_coef(group: pd.DataFrame) -> float:
        total_pop = group["Pop_all_ages"].sum()

        if total_pop == 0:
            return 0.0

        # If no gender specified (safety) -> use all genders
        if not selected_genders:
            num = group["Pop_age_range"].sum()
        else:
            num = group.loc[
                group["Gjinia"].isin(selected_genders),
                "Pop_age_range"
            ].sum()

        if num <= 0:
            return 0.0

        return float(num) / float(total_pop)

    coef_by_komuna = df.groupby("Komuna").apply(calc_coef)
    return coef_by_komuna

def compute_age_os_share(df_ga: pd.DataFrame,
                         age_cols,
                         selected_genders,
                         global_min_age: int,
                         global_max_age: int | None,
                         os_min_age: int,
                         os_max_age: int) -> pd.Series:
    """
    Kthen për secilën Komunë:
    share_OS(komuna) = Pop( OS_min–OS_max ) / Pop( global_min–global_max )

    – përdor vetëm gjinitë e zgjedhura te `selected_genders`
    – nëse s'ka popullsi në intervalin global -> 0
    """

    df = df_ga.copy()

    # Sorto kolonat e moshës
    age_cols_sorted = sorted(age_cols, key=lambda x: int(str(x)))
    max_available_age = int(str(age_cols_sorted[-1]))

    # Global max nëse është None
    if global_max_age is None:
        global_max_age = max_available_age

    # Konverto në int
    global_min_age = int(global_min_age)
    global_max_age = int(global_max_age)
    os_min_age = int(os_min_age)
    os_max_age = int(os_max_age)

    # Kufizo OS brenda intervalit global
    os_min_age = max(os_min_age, global_min_age)
    os_max_age = min(os_max_age, global_max_age)

    if os_min_age > os_max_age:
        # Interval OS bosh -> gjithmonë 0
        return pd.Series(0.0, index=df["Komuna"].unique())

    global_cols = [
        c for c in age_cols_sorted
        if global_min_age <= int(str(c)) <= global_max_age
    ]
    os_cols = [
        c for c in age_cols_sorted
        if os_min_age <= int(str(c)) <= os_max_age
    ]

    df["Pop_global"] = df[global_cols].sum(axis=1)
    df["Pop_os"] = df[os_cols].sum(axis=1)

    def agg(group: pd.DataFrame) -> float:
        if selected_genders:
            group = group[group["Gjinia"].isin(selected_genders)]

        pop_global = group["Pop_global"].sum()
        if pop_global <= 0:
            return 0.0

        pop_os = group["Pop_os"].sum()
        if pop_os <= 0:
            return 0.0

        return float(pop_os) / float(pop_global)

    share_by_komuna = df.groupby("Komuna").apply(agg)
    return share_by_komuna

def controlled_rounding(values: np.ndarray,
                        total_n: int,
                        seed: int = 42) -> np.ndarray:
    """
    Controlled rounding:
    - Start from float allocations
    - Floor all
    - Distribute remaining units based on fractional parts (probabilistic)
    - Sum-preserving & reproducible via seed
    """
    vals = np.asarray(values, dtype=float)
    if len(vals) == 0:
        return vals.astype(int)

    floors = np.floor(vals).astype(int)
    diff = int(total_n - floors.sum())

    if diff == 0:
        return floors

    fracs = vals - floors

    rng = np.random.default_rng(seed)

    if diff > 0:
        # Distribute +1 to 'diff' positions, weighted by fractional parts
        if fracs.sum() == 0:
            # if all fracs = 0, choose random indices uniformly
            indices = rng.choice(len(vals), size=diff, replace=False)
        else:
            probs = fracs / fracs.sum()
            indices = rng.choice(len(vals), size=diff, replace=False, p=probs)
        floors[indices] += 1

    elif diff < 0:
        # Too many after floor (rare, due to numeric issues).
        # Remove -diff units from smallest fractional parts.
        diff = -diff
        # Positions with smallest fracs get -1
        order = np.argsort(fracs)  # ascending
        indices = order[:diff]
        floors[indices] -= 1

    # Safety adjust if still off by 1 from numeric edge cases
    final_diff = int(total_n - floors.sum())
    if final_diff != 0 and len(vals) > 0:
        idx = rng.integers(0, len(vals))
        floors[idx] += final_diff

    return floors
   
# Helper to identify strata belonging to each oversample variable
def mask_for_oversample(grouped, variable, params):

    if variable == "Komuna":
        return grouped[base_col] == params["value"]

    if variable == "Regjion":
        return grouped[base_col] == params["value"]

    if variable == "Vendbanimi":
        return grouped["Sub"].str.endswith(params["value"])

    if variable == "Etnia":
        return grouped["Sub"].str.startswith(params["value"])

    if variable == "Gjinia":
        if "Gjinia" in grouped.columns:
            return grouped["Gjinia"] == params["value"]
        else:
            return pd.Series(False, index=grouped.index)

    if variable == "Mosha":
        # "OS" është segmenti i krijuar më herët
        if "AgeSeg" in grouped.columns:
            return grouped["AgeSeg"] == "OS"
        return pd.Series(False, index=grouped.index)

    return pd.Series(False, index=grouped.index)

def fix_minimum_allocations(
    pivot: pd.DataFrame,
    df_eth: pd.DataFrame,
    region_map: dict,
    strata_col: list,
    majority: dict,
    selected_ethnicity: list,
    min_total: int = 3,
    min_eth: int = 3,      # threshold for removing (total eth < 3)
    min_vb: int = 2        # not used for ethnicity removal now, only for settlement logic
) -> pd.DataFrame:

    pivot_fixed = pivot.copy()
    # Përdor vetëm komunat reale, jo rreshtin "Total"
    municipalities = [m for m in pivot_fixed.index if m != "Total"]

    # store initial totals for receiver limit
    initial_total = pivot_fixed["Total"].copy()

    has_ethnicity = any(
        c.startswith(("Shqiptar", "Serb", "Tjerë"))
        for c in pivot_fixed.columns
    )

    if not has_ethnicity:                        
        # ======================================================
        # 0B: Llogarit majority ethnicity për TË GJITHA komunat
        # ======================================================
        majority_all = {}

        for kom in municipalities:
            dfk = df_eth[df_eth["Komuna"] == kom]
            dfk = dfk[dfk["Etnia"] != "Total"]

            if dfk.empty:
                majority_all[kom] = None
                continue

            pop_by_eth = dfk.groupby("Etnia")["Pop_base"].sum()

            if pop_by_eth.sum() == 0:
                majority_all[kom] = None
            else:
                majority_all[kom] = pop_by_eth.idxmax()

        # ======================================================
        # 0C: Gjej komunat me Total = 1
        # ======================================================
        mun_total_1 = [
            kom for kom in municipalities
            if int(pivot_fixed.at[kom, "Total"]) == 1
        ]

        # ======================================================
        # 0D: RIALOKO KOMUNAT KU majority ≠ selected_ethnicity
        # ======================================================
        for kom in mun_total_1:
            maj = majority_all.get(kom)

            # Përndryshe → duhet ta largojmë 1 intervistë
            reg = region_map.get(kom, None)

            # Donatorët brenda rajonit
            donors = [
            d for d in municipalities
            if d != kom
            and str(majority_all.get(d)).strip().lower() == str(selected_ethnicity[0]).strip().lower()
            and region_map.get(d, None) == reg
            and pivot_fixed.at[d, "Total"] > min_total
            ]

            # Nëse s’ka në rajon → kombëtarisht
            if not donors:
                donors = [
                    d for d in municipalities
                    if d != kom
                    and str(majority_all.get(d)).strip().lower() == str(selected_ethnicity[0]).strip().lower()
                    and pivot_fixed.at[d, "Total"] > min_total
                ]

            if not donors:
                continue

            donor = donors[0]

            # REMOVE interview from KOM
            sub_remove = pivot_fixed.loc[kom, ["Urban", "Rural"]].idxmax()
            pivot_fixed.at[kom, sub_remove] -= 1

            # ADD interview to DONOR
            sub_add = pivot_fixed.loc[donor, ["Urban", "Rural"]].idxmax()
            pivot_fixed.at[donor, sub_add] += 1

            # Recalculate totals consistently
            pivot_fixed["Total"] = pivot_fixed["Urban"] + pivot_fixed["Rural"]

        # ======================================================
        # 0E: Final totals + remove rows with Total=0
        # ======================================================
        pivot_fixed["Total"] = pivot_fixed[strata_col].sum(axis=1)
        pivot_fixed = pivot_fixed[pivot_fixed["Total"] != 0]
        pivot_fixed.loc["Total"] = pivot_fixed.sum(numeric_only=True)

        return pivot_fixed
        

    ##############################################################
    # Detect ethnicity structure automatically
    ##############################################################

    eth_basic = ["Shqiptar", "Serb", "Tjerë"]

    eth_groups = {}
    for eth in eth_basic:
        cols_for_eth = [c for c in pivot_fixed.columns if str(c).startswith(eth)]
        if cols_for_eth:
            eth_groups[eth] = cols_for_eth

    if eth_groups:
        eth_structure = "eth_dynamic"
    else:
        eth_structure = "none"

    # ------------------------------------------------------
    # If there is no usable ethnicity structure, exit early
    # ------------------------------------------------------
    if eth_structure == "none":
        # Only ensure minimum TOTAL per municipality
        for kom in municipalities:
            deficit = min_total - pivot_fixed.at[kom, "Total"]
            if deficit > 0:
                # find donors in same region
                region = region_map.get(kom, None)
                donors = [d for d in municipalities if d != kom and region_map.get(d, None) == region]

                if not donors:
                    donors = [d for d in municipalities if d != kom]

                for d in donors:
                    if deficit == 0:
                        break
                    if pivot_fixed.at[d, "Total"] > min_total:
                        pivot_fixed.at[d, "Total"] -= 1
                        pivot_fixed.at[kom, "Total"] += 1
                        deficit -= 1

        pivot_fixed.loc["Total"] = pivot_fixed.sum(numeric_only=True)
        return pivot_fixed

    # allowed matrix (columns that existed initially)
    allowed = (pivot > 0)

    # helper — find receivers for ethnicity
    def receiver_candidates(eth, col, donor_kom):

        receivers = []
        break_reasons = {}   # store reasons why each r was blocked (nëse të duhet për debug)

        for r in pivot_fixed.index:

            # Skip donor itself
            if r == donor_kom:
                break_reasons[r] = "same municipality - skip"
                continue

            # RULE 1: majority restriction
            if eth in ["Serb", "Shqiptar"]:
                if majority.get(r) != eth:
                    break_reasons[r] = f"blocked: majority[{r}] = {majority.get(r)}, required = {eth}"
                    continue

            # RULE 2: allowed matrix (ethnicity–settlement allowed or not)
            if col not in allowed.columns:
                break_reasons[r] = f"blocked: column {col} not in allowed"
                continue

            if not allowed.at[r, col]:
                break_reasons[r] = f"blocked: allowed[{r},{col}] = False"
                continue

            # RULE 3: donor limit (+3 rule)
            if pivot_fixed.at[r, "Total"] >= initial_total[r] + 3:
                break_reasons[r] = (
                    f"blocked: total {pivot_fixed.at[r,'Total']} >= initial_total + 3 "
                    f"({initial_total[r]} + 3 = {initial_total[r] + 3})"
                )
                continue

            receivers.append(r)

        # Apply regional priority
        region_of = {m: region_map.get(m, None) for m in pivot_fixed.index}
        donor_region = region_of.get(donor_kom, None)

        if donor_region is not None:
            in_region = [rcv for rcv in receivers if region_of.get(rcv, None) == donor_region]
        else:
            in_region = []

        if in_region:
            return in_region

        return receivers

    # -----------------------------------------------------
    # ETHNIC REALLOCATION (core logic)
    # -----------------------------------------------------
    # We keep Urban/Rural separately where applicable, but remove all
    # units for an ethnicity in a municipality if total < min_eth
    # -----------------------------------------------------

    for kom in municipalities:

        for eth, cols in eth_groups.items():

            # total across all columns for this ethnicity (Urban/Rural or single)
            total_eth = sum(pivot_fixed.at[kom, c] for c in cols if c in pivot_fixed.columns)

            # OK if >= min_eth
            if total_eth >= min_eth:
                continue

            # nothing to remove if 0
            if total_eth == 0:
                continue

            # number of units to remove = all units
            units_to_move = total_eth  # (përdoret konceptualisht, nuk është i domosdoshëm në while)

            # move across columns (p.sh. fillimisht Urban pastaj Rural)
            for col in cols:
                if col not in pivot_fixed.columns:
                    continue

                while pivot_fixed.at[kom, col] > 0:

                    # find receivers
                    recv_list = receiver_candidates(eth, col, kom)

                    if not recv_list:
                        break

                    recv = recv_list[0]

                    # SAFETY: donor cannot go negative in this ethnicity column
                    if pivot_fixed.at[kom, col] <= 0:
                        break

                    # SAFETY: donor municipality total cannot go below min_total
                    if pivot_fixed.at[kom, "Total"] <= min_total:
                        break

                    # SAFETY: receiver cannot exceed initial_total + 3
                    if pivot_fixed.at[recv, "Total"] >= initial_total[recv] + 3:
                        break

                    # transfer 1 unit FROM kom TO recv
                    pivot_fixed.at[kom, col] -= 1
                    pivot_fixed.at[recv, col] += 1

                    pivot_fixed.at[kom, "Total"] -= 1
                    pivot_fixed.at[recv, "Total"] += 1

        # ---------------------------------------------------------
        # FINAL PASS: Fix ethnicity totals where combined total = 1
        # (bëhet si për eth_only ashtu edhe për eth_settlement)
        # ---------------------------------------------------------

        eth_groups_final = eth_groups  # përdor grupet e detektuara nga eth_structure

        for kom in municipalities:

            for eth, cols in eth_groups_final.items():

                # Combined total of this ethnicity in this municipality
                total_eth = sum(
                    pivot_fixed.at[kom, c]
                    for c in cols
                    if c in pivot_fixed.columns
                )

                # Only fix the EXACT 1-case
                if total_eth != 1:
                    continue

                # -------------------------------------------------
                # 1) Gjej donor të vlefshëm
                # -------------------------------------------------
                donors = []
                for d in municipalities:
                    if d == kom:
                        continue

                    # RULE: Majority restriction (vetëm kur zbatohet)
                    if eth in ["Serb", "Shqiptar"] and majority.get(d) != eth:
                        continue

                    donor_eth_total = sum(
                        pivot_fixed.at[d, c]
                        for c in cols
                        if c in pivot_fixed.columns
                    )
                    if donor_eth_total <= min_eth:
                        continue

                    # Donor duhet të ketë njësi për dhënë
                    if all(
                        pivot_fixed.at[d, c] <= 0
                        for c in cols
                        if c in pivot_fixed.columns
                    ):
                        continue

                    # Donor total nuk mund të bjerë nën min_total
                    if pivot_fixed.at[d, "Total"] <= min_total:
                        continue

                    donors.append(d)

                if not donors:
                    continue

                # Prefer same region donor
                region_kom = region_map.get(kom)
                donors_in_region = [
                    d for d in donors
                    if region_map.get(d) == region_kom
                ]

                donor = donors_in_region[0] if donors_in_region else donors[0]

                # -------------------------------------------------
                # 2) Zgjedh kolonën ku ta shtoj (një nga cols)
                # -------------------------------------------------
                # Prioritet sipas allowed matrix, nëse ekziston
                target_col = None
                for c in cols:
                    if c in allowed.columns and allowed.at[kom, c]:
                        target_col = c
                        break

                if target_col is None:
                    # fallback: merr kolonën e parë ekzistuese
                    for c in cols:
                        if c in pivot_fixed.columns:
                            target_col = c
                            break

                if target_col is None:
                    continue

                # -------------------------------------------------
                # 3) Zgjedh kolonën nga ku do të heqim te donori
                # -------------------------------------------------
                donor_col = None
                for c in cols:
                    if c in pivot_fixed.columns and pivot_fixed.at[donor, c] > 0:
                        donor_col = c
                        break

                if donor_col is None:
                    continue

                # -------------------------------------------------
                # 4) Transferimi i njësisë
                # -------------------------------------------------
                pivot_fixed.at[kom, target_col] += 1
                pivot_fixed.at[donor, donor_col] -= 1

                pivot_fixed.at[kom, "Total"] += 1
                pivot_fixed.at[donor, "Total"] -= 1

    # -----------------------------------------------------
    # RECOMPUTE TOTALS
    # -----------------------------------------------------

    pivot_fixed["Total"] = pivot_fixed[strata_col].sum(axis=1)

    # Ribëj totalin e fundit
    pivot_fixed.loc["Total"] = pivot_fixed.sum(numeric_only=True)

    return pivot_fixed

@st.cache_data
def load_psu_data(path: str) -> pd.DataFrame:
    df = pd.read_excel(path)
    # Normalizim minimal
    df["Komuna"] = df["Komuna"].astype(str).str.strip()
    df["Vendbanimi"] = df["Vendbanimi"].astype(str).str.strip()
    df["Fshati/Qyteti"] = df["Fshati/Qyteti"].astype(str).str.strip()
    df["Quadrant"] = df["Quadrant"].astype(str).str.strip()

    # Etnitë kryesore
    for col in ["Shqiptar", "Serb"]:
        if col not in df.columns:
            df[col] = 0.0
        df[col] = df[col].fillna(0).astype(float)

    other_cols = [
        "Boshnjak", "Turk", "Rom", "Ashkali", "Egjiptian",
        "Goran", "Të tjerë", "Preferoj të mos përgjigjem"
    ]
    for col in other_cols:
        if col not in df.columns:
            df[col] = 0.0
        df[col] = df[col].fillna(0).astype(float)

    df["Tjeter_pop"] = df[other_cols].sum(axis=1)

    return df

def compute_num_psu(total_interviews: int, k: int):
    """
    Rregulli yt:
    - q = T // k, r = T % k
    - nëse r == 0 → PSU të plota
    - nëse r <= k/2 → nuk shtohet PSU, leftover = r (shpërndahet te PSU-të më të mëdha)
    - nëse r > k/2 → shtohet një PSU shtesë me madhësi r
    """
    if total_interviews <= 0:
        return 0, 0, 0

    q = total_interviews // k
    r = total_interviews % k
    half_k = k / 2.0

    if r == 0:
        return q, 0, 0

    if r <= half_k:
        return q, r, 0
    else:
        return q + 1, 0, r


def select_psus_for_municipality(
    komuna: str,
    total_interviews: int,
    df_psu_mun: pd.DataFrame,
    k: int,
    required_ethnicities: list[str]
) -> pd.DataFrame:
    """
    Zgjedh PSU-të për një komunë:
    - garanton sa më shumë që të jetë e mundur përfaqësim të quadrant-eve
    - tenton të ketë të paktën një PSU ku ekziston çdo etni e kërkuar
    - shpërndan intervistat sipas rregullit të k/num_psu/leftover
    """

    df = df_psu_mun.copy()
    if df.empty or total_interviews <= 0:
        return pd.DataFrame()

    num_psu, leftover, extra_psu_size = compute_num_psu(total_interviews, k)

    if num_psu == 0:
        return pd.DataFrame()

    df["PopFilt"] = df.apply(
    lambda r: compute_filtered_pop_for_psu_row(
        r,
        age_min=min_age,
        age_max=max_age,
        gender_selected=gender_selected,
        eth_filter=eth_filter
    ),
    axis=1)

    ALL_ETH_COLS = [
        "Shqiptar", "Serb", "Boshnjak", "Turk", "Rom",
        "Ashkali", "Egjiptian", "Goran", "Të tjerë",
        "Preferoj të mos përgjigjem"
    ]

    OTHER_ETH_COLS = [e for e in ALL_ETH_COLS if e not in ["Shqiptar", "Serb"]]

    def compute_ethnic_pop_filtered(row):
        eth_total = sum(row.get(c, 0) for c in ALL_ETH_COLS)
        if eth_total <= 0:
            return pd.Series({
                "Shqiptar_pop": 0.0,
                "Serb_pop": 0.0,
                "Tjeter_pop": 0.0
            })

        shq = row.get("Shqiptar", 0) / eth_total
        ser = row.get("Serb", 0) / eth_total
        tjr = sum(row.get(c, 0) for c in OTHER_ETH_COLS) / eth_total

        return pd.Series({
            "Shqiptar_pop": row["PopFilt"] * shq,
            "Serb_pop":     row["PopFilt"] * ser,
            "Tjeter_pop":   row["PopFilt"] * tjr
        })

    # Remove any previous duplicate Tjeter_pop column
    if "Tjeter_pop" in df.columns:
        df = df.drop(columns=["Tjeter_pop"])

    eth_cols_df = df.apply(compute_ethnic_pop_filtered, axis=1)
    df = pd.concat([df, eth_cols_df], axis=1)

    # --------------------------
    # 1) Përfaqësimi i quadrant-eve
    # --------------------------
    df = df.sort_values("PopFilt", ascending=False)
    quads = df["Quadrant"].dropna().unique().tolist()

    selected_idx = []

    # a) nëse kemi mjaftueshëm PSU për të gjithë quadrant-et
    if num_psu >= len(quads):
        for q in quads:
            cand = df[df["Quadrant"] == q]
            if not cand.empty:
                selected_idx.append(cand.index[0])

        # plotëso numrin e PSU-ve me PSU-të më të mëdha të mbetura
        remaining_needed = num_psu - len(selected_idx)
        if remaining_needed > 0:
            remaining = df.drop(index=selected_idx)
            extra = remaining.head(remaining_needed).index.tolist()
            selected_idx.extend(extra)

    else:
        # b) pak PSU → zgjedhim si fillim PSU-në më të madhe brenda çdo quadrant-i
        top_per_quad = (
            df.sort_values("PopFilt", ascending=False)
            .groupby("Quadrant", group_keys=False)
            .head(1)
        )

        # tani rendisim këto top-PSU sipas popullsisë, dhe marrim vetëm aq sa na duhen
        top_per_quad = top_per_quad.sort_values("PopFilt", ascending=False)

        selected_idx = top_per_quad.head(num_psu).index.tolist()

        # nëse akoma s'e kemi arritur numrin, plotëso me më të mëdhenjtë
        if len(selected_idx) < num_psu:
            remaining = df.drop(index=selected_idx)
            extra = remaining.head(num_psu - len(selected_idx)).index.tolist()
            selected_idx.extend(extra)

            
    selected_idx = list(dict.fromkeys(selected_idx))  # heq duplikate duke ruajtur rendin
    selected = df.loc[selected_idx].copy()

    # --------------------------
    # 2) Siguro prezencën e etnive të kërkuara
    # --------------------------
    def has_eth(selected_df, eth: str) -> bool:
        if eth == "Shqiptar":
            return (selected_df["Shqiptar_pop"] > 0).any()
        if eth == "Serb":
            return (selected_df["Serb_pop"] > 0).any()
        if eth == "Tjerë":
            return (selected_df["Tjeter_pop"] > 0).any()
        return False

    for eth in required_ethnicities:
        if has_eth(selected, eth):
            continue

        # gjej PSU jashtë të zgjedhurave që ka këtë etni
        if eth == "Shqiptar":
            cand_eth = df[(df["Shqiptar_pop"] > 0) & (~df.index.isin(selected.index))]
        elif eth == "Serb":
            cand_eth = df[(df["Serb_pop"] > 0) & (~df.index.isin(selected.index))]
        elif eth == "Tjerë":
            cand_eth = df[(df["Tjeter_pop"] > 0) & (~df.index.isin(selected.index))]
        else:
            continue

        if cand_eth.empty:
            # nuk ka PSU me këtë etni në këtë komunë
            continue

        # marrim kandidatin më të madh
        new_psu = cand_eth.iloc[0]

        # gjej një PSU për t'u zëvendësuar që nuk është i vetmi në quadrant-in e vet
        removed_idx = None
        for idx, row in selected.sort_values("PopFilt", ascending=False).iterrows():
            q = row["Quadrant"]
            # a ka PSU të tjera në të njëjtin quadrant në 'selected'?
            if (selected["Quadrant"] == q).sum() > 1:
                removed_idx = idx
                break

        if removed_idx is None:
            # nuk mund të bëjmë swap pa prishur quadrant-et → si fallback mund ta shtojmë
            # por për të mos prishur logjikën e numrit të PSU-ve, aktualisht e anashkalojmë
            continue

        # bëjmë zëvendësimin
        selected = selected.drop(index=removed_idx)
        selected = pd.concat([selected, new_psu.to_frame().T])

    # --------------------------
    # 3) Shpërndarja e intervistave te PSU-të
    # --------------------------
    selected = selected.sort_values("PopFilt", ascending=False).reset_index(drop=True)

    if extra_psu_size > 0:
        # p.sh. 46 anketa, k=8 → 6 PSU (5*8 + 6)
        base_sizes = [k] * (num_psu - 1) + [extra_psu_size]
    else:
        # p.sh. 42 anketa, k=8 → 5 PSU (5*8) + leftover=2 → shpërndajmë 2
        base_sizes = [k] * num_psu

        if leftover > 0:
            N = len(base_sizes)

            # ------------------------------
            # RULE A: leftover <= N
            # ------------------------------
            if leftover <= N:
                for i in range(leftover):
                    base_sizes[i] += 1

            else:
                # ------------------------------
                # RULE B: leftover > N
                # ------------------------------

                # STEP 1: give 1 interview to each PSU
                for i in range(N):
                    base_sizes[i] += 1

                L = leftover - N  # remaining interviews

                # STEP 2: distribute remaining leftover round-robin
                idx = 0
                while L > 0:
                    base_sizes[idx] += 1
                    L -= 1
                    idx = (idx + 1) % N

    selected["Intervista"] = base_sizes[: len(selected)]

    # shtojmë info etnish (num popullsie në atë PSU)
    selected["Shqiptar_pop"] = selected["Shqiptar_pop"].astype(float)
    selected["Serb_pop"] = selected["Serb_pop"].astype(float)
    selected["Tjeter_pop"] = selected["Tjeter_pop"].astype(float)

    selected["Komuna"] = komuna

    return selected[
        [
            "Komuna",
            "Fshati/Qyteti",
            "Vendbanimi",
            "Intervista"
        ]
    ]

@st.cache_data(show_spinner="Duke ngarkuar listën e kontakteve...")
def load_citizens_database():
    # Merr kredencialet nga st.secrets
    gcp_info = st.secrets["gcp_service_account"]
    
    # Deklaro scope të qartë për Google Sheets
    scopes = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive"]
    
    # Krijo kredencialet me scope
    credentials = Credentials.from_service_account_info(gcp_info, scopes=scopes)
    
    # Autorizo me gspread
    gc = gspread.authorize(credentials)
    
    # Hap dokumentin dhe worksheet-in
    sheet = gc.open("Databaza e kontakteve të qytetarëve").worksheet("Sheet1")
    df = pd.DataFrame(sheet.get_all_records())

    return df

def compute_psu_table_for_all_municipalities(
    pivot: pd.DataFrame,
    df_psu: pd.DataFrame,
    k: int,
    eth_filter: list[str],
    settlement_filter: list[str]
) -> pd.DataFrame:
    """
    Gjeneron tabelën finale të PSU-ve për të gjitha komunat.
    - Urban = 1 rresht me total Urban intervista
    - Rural = përdor select_psus_for_municipality()
    """

    def extract_urban_interviews(pivot_row):
        urban_cols = [c for c in pivot_row.index if "Urban" in str(c)]
        return int(pivot_row[urban_cols].sum()) if urban_cols else 0

    # Gjej kolonat e etnisë në pivot
    eth_cols_map = {
        eth: [
            c for c in pivot.columns
            if c != "Total" and str(c).startswith(eth)
        ]
        for eth in ["Shqiptar", "Serb", "Tjerë"]
    }

    all_rows = []

    for kom in pivot.index:
        if kom == "Total":
            continue

        # pikënisja
        pivot_row = pivot.loc[kom]
        total_interviews = int(pivot_row["Total"])
        if total_interviews <= 0:
            continue

        # Llogarit Urban dhe Rural
        urban_int = extract_urban_interviews(pivot_row)
        rural_int = total_interviews - urban_int

        df_mun = df_psu[df_psu["Komuna"] == kom].copy()

        if df_mun.empty:
            continue

        # Filtrim sipas vendbanimit të zgjedhur (nëse e përdor në app)
        if settlement_filter:
            df_mun = df_mun[df_mun["Vendbanimi"].isin(settlement_filter)]

        # ===========================
        # 1) URBAN PSU (një rresht)
        # ===========================
        if urban_int >= 0:
            df_mun_urban = df_mun[df_mun["Vendbanimi"] == "Urban"]

            if not df_mun_urban.empty:
                # There is always exactly 1 Urban row per municipality
                best_urban = df_mun_urban.iloc[0].copy()

                # Compute PopFilt for Urban row
                best_urban["PopFilt"] = compute_filtered_pop_for_psu_row(
                    best_urban,
                    age_min=min_age,
                    age_max=max_age,
                    gender_selected=gender_selected,
                    eth_filter=eth_filter
                )


                row_urban = pd.DataFrame([{
                    "Komuna": kom,
                    "Fshati/Qyteti": best_urban["Fshati/Qyteti"],
                    "Vendbanimi": "Urban",
                    "Quadrant": "-",
                    "PopFilt": best_urban["PopFilt"],
                    "Intervista": urban_int,
                    "Shqiptar_pop": best_urban.get("Shqiptar_pop", 0),
                    "Serb_pop": best_urban.get("Serb_pop", 0),
                    "Tjeter_pop": best_urban.get("Tjeter_pop", 0)
                }])

                all_rows.append(row_urban)


        # ===========================
        # 2) RURAL PSU
        # ===========================

        # Gjej cilat etni kanë mostra > 0 në këtë komunë
        required_eth = []
        for eth, cols in eth_cols_map.items():
            if eth not in eth_filter:
                continue
            if not cols:
                continue
            if int(pivot.loc[kom, cols].sum()) > 0:
                required_eth.append(eth)

        if rural_int > 0:
            df_mun_rural = df_mun[df_mun["Vendbanimi"] == "Rural"]

            psu_rural = select_psus_for_municipality(
                komuna=kom,
                total_interviews=rural_int,
                df_psu_mun=df_mun_rural,
                k=k,
                required_ethnicities=required_eth
            )

            # FALLBACK: if selection fails (empty), allocate all interviews to largest Rural PSU
            if psu_rural.empty:

                # Compute PopFilt for rural PSU rows BEFORE sorting
                df_mun_rural = df_mun_rural.copy()
                df_mun_rural["PopFilt"] = df_mun_rural.apply(
                    lambda r: compute_filtered_pop_for_psu_row(
                        r,
                        age_min=min_age,
                        age_max=max_age,
                        gender_selected=gender_selected,
                        eth_filter=eth_filter
                    ),
                    axis=1
                )

                # If still empty (no valid PopFilt), return at least one PSU
                if df_mun_rural["PopFilt"].sum() == 0:
                    # fallback: pick first rural PSU
                    fallback_row = df_mun_rural.head(1)
                    fallback_row = fallback_row.assign(Intervista=rural_int)[
                        ["Komuna", "Fshati/Qyteti", "Vendbanimi", "Intervista"]
                    ]
                    all_rows.append(fallback_row)
                else:
                    # sort by PopFilt
                    fallback = df_mun_rural.sort_values("PopFilt", ascending=False).head(1)
                    fallback = fallback.assign(Intervista=rural_int)[
                        ["Komuna", "Fshati/Qyteti", "Vendbanimi", "Intervista"]
                    ]
                    all_rows.append(fallback)

            else:
                all_rows.append(psu_rural)


    # ===========================
    # 3) Bashkimi final
    # ===========================
    if not all_rows:
        return pd.DataFrame()

    final_psu = pd.concat(all_rows, ignore_index=True)
    return final_psu[
        [
            "Komuna",
            "Fshati/Qyteti",
            "Vendbanimi",
            "Intervista"
        ]
    ]

def extract_urban_interviews(pivot_row):
    urban_cols = [c for c in pivot_row.index if "Urban" in str(c)]
    return int(pivot_row[urban_cols].sum()) if urban_cols else 0

def compute_filtered_pop_for_psu_row(
    psu_row: pd.Series,
    age_min: int,
    age_max: int | None,
    gender_selected: list[str],
    eth_filter: list[str]
) -> float:
    """
    Compute population for one PSU using demographic filters.
    Fully safe: no division by zero, respects gender/age/eth filters.
    """

    import re

    # -------------------------------------------------------
    # 1. Handle age max
    # -------------------------------------------------------
    if age_max is None:
        age_max = 120  # large value = include all

    # -------------------------------------------------------
    # 2. Identify all PSU age group columns
    # -------------------------------------------------------
    age_cols = []
    for col in psu_row.index:
        name = str(col).strip()
        # match formats like '0-4', '5-9', '65+', etc.
        if re.fullmatch(r"\d+\-\d+", name) or re.fullmatch(r"\d+\+", name):
            age_cols.append(col)

    # Helper to decode age range
    def group_range(col_name):
        if "+" in col_name:
            base = int(col_name.replace("+", ""))
            return (base, 200)
        lo, hi = col_name.split("-")
        return (int(lo), int(hi))

    # -------------------------------------------------------
    # 3. AGE FILTER (fractional overlap)
    # -------------------------------------------------------
    pop_age = 0
    for col in age_cols:
        lo, hi = group_range(col)
        group_pop = psu_row[col]

        if group_pop <= 0:
            continue

        # overlap between age group and filter
        overlap_lo = max(lo, age_min)
        overlap_hi = min(hi, age_max)

        if overlap_lo > overlap_hi:
            continue  # no overlap

        group_size = hi - lo + 1
        overlap_size = overlap_hi - overlap_lo + 1

        fraction = overlap_size / group_size

        pop_age += group_pop * fraction

    # If no ages match → return 0
    if pop_age <= 0:
        return 0

    # -------------------------------------------------------
    # 4. GENDER FILTER (SAFE)
    # -------------------------------------------------------
    male = psu_row.get("Meshkuj", 0)
    female = psu_row.get("Femra", 0)

    if "Meshkuj" not in gender_selected:
        male = 0
    if "Femra" not in gender_selected:
        female = 0

    total_gender_pop = male + female

    if total_gender_pop <= 0:
        return 0

    # Proportion from allowed genders
    gender_fraction = total_gender_pop / (psu_row.get("Meshkuj", 0) + psu_row.get("Femra", 0)
                                          if (psu_row.get("Meshkuj", 0) + psu_row.get("Femra", 0)) > 0
                                          else total_gender_pop)

    # -------------------------------------------------------
    # 5. ETHNICITY FILTER (SAFE)
    # -------------------------------------------------------
    eth_total = 0
    eth_selected = 0

    all_ethnic_cols = [
        "Shqiptar", "Serb", "Boshnjak", "Turk", "Rom",
        "Ashkali", "Egjiptian", "Goran", "Të tjerë",
        "Preferoj të mos përgjigjem"
    ]

    for eth in all_ethnic_cols:
        eth_total += psu_row.get(eth, 0)

    for eth in eth_filter:
        eth_selected += psu_row.get(eth, 0)

    # SAFE division
    if eth_total > 0:
        eth_fraction = eth_selected / eth_total
    else:
        eth_fraction = 0

    if eth_fraction <= 0:
        return 0

    # -------------------------------------------------------
    # 6. Combine all three filters (age × gender × ethnicity)
    # -------------------------------------------------------
    final_pop = pop_age * (total_gender_pop / (male + female) if (male + female) > 0 else 1)
    final_pop *= eth_fraction

    return max(final_pop, 0)

def filter_contacts(
    df_contacts: pd.DataFrame,
    *,
    komuna_filter: list[str],
    gender_selected: list[str],
    min_age: int,
    max_age: int | None,
    settlement_filter: list[str]
) -> pd.DataFrame:

    df = df_contacts.copy()

    # --------------------------------------------------
    # 1. Mandatory non-null checks
    # --------------------------------------------------
    df = df[
        df["Emri dhe mbiemri"].notna() &
        df["Numri i telefonit"].notna()
    ]

    df = df[
        (df["Emri dhe mbiemri"] != "0") &
        (df["Numri i telefonit"] != "0")
    ]

    # --------------------------------------------------
    # 3. Municipality filter (from sidebar filters)
    # --------------------------------------------------
    if komuna_filter:
        df = df[df["Komuna"].isin(komuna_filter)]

    # --------------------------------------------------
    # 4. Gender filter
    # --------------------------------------------------
    if gender_selected:

        gender_map = {
        "Meshkuj": "Mashkull",
        "Femra": "Femër"
        }

        mapped_selected = [gender_map.get(g, g) for g in gender_selected]

        df = df[df["Gjinia"].isin(mapped_selected)]

    # --------------------------------------------------
    # 5. Age filter
    # --------------------------------------------------
    df["Mosha"] = pd.to_numeric(df["Mosha"], errors="coerce")

    if max_age is None:
        df = df[df["Mosha"] >= min_age]
    else:
        df = df[(df["Mosha"] >= min_age) & (df["Mosha"] <= max_age)]

    # --------------------------------------------------
    # 6. Settlement filter
    # --------------------------------------------------
    if settlement_filter:
        df = df[df["Vendbanimi"].isin(settlement_filter)]
    
    return df.reset_index(drop=True)

def build_contact_list(
    df_contacts: pd.DataFrame,
    pivot: pd.DataFrame,
    *,
    reserve_mode: str,
    reserve_percentage: int | None,
    reserve_ratio: int | None,
    seed: int = 42
) -> pd.DataFrame:

    rng = np.random.default_rng(seed)
    rows = []

    for komuna in pivot.index:
        if komuna == "Total":
            continue

        for sub in pivot.columns:
            if sub == "Total":
                continue

            interviews = int(pivot.at[komuna, sub])
            if interviews <= 0:
                continue

            # ----------------------------
            # Decode secondary stratum
            # ----------------------------
            eth = None
            settlement = None

            if " - " in sub:
                eth, settlement = sub.split(" - ")
            elif sub in ["Urban", "Rural"]:
                settlement = sub
            elif sub in ["Shqiptar", "Serb", "Tjerë"]:
                eth = sub

            # ----------------------------
            # Filter contacts
            # ----------------------------
            dfk = df_contacts[df_contacts["Komuna"] == komuna]

            if eth:
                dfk = dfk[dfk["Etnia"] == eth]

            if settlement:
                dfk = dfk[dfk["Vendbanimi"] == settlement]

            if dfk.empty:
                continue

            # ----------------------------
            # Required contacts
            # ----------------------------
            if reserve_mode == "Proporcion":
                required = interviews * reserve_ratio
            else:
                required = interviews + int(
                    np.ceil(interviews * reserve_percentage / 100)
                )

            # ----------------------------
            # Sample
            # ----------------------------
            if len(dfk) <= required:
                sample = dfk.copy()
            else:
                sample = dfk.sample(
                    n=required,
                    replace=False,
                    random_state=seed
                )

            sample["Komuna"] = komuna
            sample["Stratum"] = sub
            sample["Intervista_plan"] = interviews
            sample["Kontakt_kërkuar"] = required
            sample["Kontakt_marrë"] = len(sample)

            rows.append(sample)

    if not rows:
        return pd.DataFrame()

    return pd.concat(rows, ignore_index=True)

def add_markdown_runs(paragraph, text):
    """Adds runs with markdown formatting: bold, italic, bold+italic."""
    bold_italic = r"\*\*\*(.+?)\*\*\*"
    bold = r"\*\*(.+?)\*\*"
    italic = r"\*(.+?)\*"

    tokens = re.split(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)", text)

    for token in tokens:
        run = paragraph.add_run()
        if re.match(bold_italic, token):
            run.text = re.findall(bold_italic, token)[0]
            run.bold = True
            run.italic = True
        elif re.match(bold, token):
            run.text = re.findall(bold, token)[0]
            run.bold = True
        elif re.match(italic, token):
            run.text = re.findall(italic, token)[0]
            run.italic = True
        else:
            run.text = token


def narrative_to_word(markdown_text: str) -> bytes:
    doc = Document()

    # REMOVE double blank lines (convert them to a single paragraph)
    lines = markdown_text.split("\n")
    cleaned_lines = []
    skip_next = False

    for i in range(len(lines)):
        if i < len(lines) - 1 and lines[i].strip() == "" and lines[i+1].strip() == "":
            # skip multiple blank lines
            if not skip_next:
                cleaned_lines.append("")  # keep one blank line
            skip_next = True
        else:
            cleaned_lines.append(lines[i])
            skip_next = False

    # CONVERT TO WORD PARAGRAPHS
    for line in cleaned_lines:

        # ----- HEADINGS (keep centered/left as Word default) -----
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            continue
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            continue
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            continue

        # ----- BULLET POINTS -----
        if line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_markdown_runs(p, line.strip()[2:])
            continue

        # ----- BLANK LINE -----
        if line.strip() == "":
            continue

        # ----- NORMAL PARAGRAPH -----
        p = doc.add_paragraph()
        add_markdown_runs(p, line)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Return bytes file
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def add_markdown_runs(paragraph, text):
    """Adds runs with markdown formatting: bold, italics, bold+italic."""

    # Patterns:
    bold_italic = r"\*\*\*(.+?)\*\*\*"
    bold = r"\*\*(.+?)\*\*"
    italic = r"\*(.+?)\*"

    # Tokenize using a combined regex
    tokens = re.split(r"(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)", text)

    for token in tokens:
        run = paragraph.add_run()

        if re.match(bold_italic, token):
            run.text = re.findall(bold_italic, token)[0]
            run.bold = True
            run.italic = True

        elif re.match(bold, token):
            run.text = re.findall(bold, token)[0]
            run.bold = True

        elif re.match(italic, token):
            run.text = re.findall(italic, token)[0]
            run.italic = True

        else:
            run.text = token

def df_to_excel_bytes(df: pd.DataFrame, sheet_name: str = "Data") -> bytes:
        from io import BytesIO
        output = BytesIO()
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=True, sheet_name=sheet_name)
        return output.getvalue()

def create_download_link(file_bytes: bytes, filename: str, label: str):
    """Create full-width HTML download button (without rerun)."""
    b64 = base64.b64encode(file_bytes).decode()
    button_html = f"""<a href="data:application/octet-stream;base64,{b64}" download="{filename}" style="text-decoration:none;">
                <div style="
                    background-color:#344b77;
                    color:white;
                    text-align:center;
                    font-weight:500;
                    font-size:16px;
                    padding:10px;
                    border-radius:8px;
                    margin-top:8px;
                    width:100%;
                    box-sizing:border-box;
                    cursor:pointer;
                ">
                {label}
                </div>
            </a>
        """
    st.markdown(button_html, unsafe_allow_html=True)

def create_download_link2(file_bytes: bytes, filename: str, label: str):
    """Create full-width HTML download button (without rerun)."""
    b64 = base64.b64encode(file_bytes).decode()
    button_html = f"""<a href="data:application/octet-stream;base64,{b64}" download="{filename}" style="text-decoration:none;">
                <div style="
                    background-color:#5b8fb8;
                    color:white;
                    text-align:center;
                    font-weight:500;
                    font-size:16px;
                    padding:10px;
                    border-radius:8px;
                    margin-top:8px;
                    width:100%;
                    box-sizing:border-box;
                    cursor:pointer;
                ">
                {label}
                </div>
            </a>
        """
    st.markdown(button_html, unsafe_allow_html=True)
def compute_population_coefficients(
    df_ga,
    df_eth,
    region_map,
    gender_selected,
    min_age,
    max_age,
    eth_filter,
    settlement_filter,
    komuna_filter,
    data_collection_method
):
    """
    Funksioni FINAL profesional për llogaritjen e peshave të popullsisë.
    Punon me df_eth në format LONG (Komuna, Vendbanimi, Etnia, Pop_base)
    dhe df_ga në formatin (Komuna, Gjinia, 0..120, Gjithsej).

    UNIVERSI I POPULLSISË = popullsia e komunës min_age → max_age (default 120).
    """

    out = []

    # -------------------------------------------------------------
    # 0) Filtrimi i df_eth dhe df_ga për komunat e kërkuara
    # -------------------------------------------------------------
    df_eth_f = df_eth[df_eth["Komuna"].isin(komuna_filter)].copy()
    df_eth_f = df_eth_f[df_eth_f["Vendbanimi"].isin(settlement_filter)]

    if eth_filter:  # Nëse janë zgjedhur etni specifike
        df_eth_f = df_eth_f[df_eth_f["Etnia"].isin(eth_filter)]

    if df_eth_f.empty:
        return pd.DataFrame()

    df_ga_f = df_ga[df_ga["Komuna"].isin(komuna_filter)].copy()
    df_ga_f = df_ga_f[df_ga_f["Gjinia"].isin(gender_selected)]

    if df_ga_f.empty:
        return pd.DataFrame()

    # -------------------------------------------------------------
    # 1) Defino moshat dhe zgjedh kolonat për kalkulim
    # -------------------------------------------------------------
    age_columns = []
    for c in df_ga.columns:
        try:
            age_columns.append(int(c))
        except:
            continue

    # Nëse nuk ka max_age → max_age = 120
    if max_age is None:
        max_age = 200

    selected_age_cols = [a for a in age_columns if min_age <= a <= max_age]

    # -------------------------------------------------------------
    # 2) Derived population per Komuna = Σ(mosha min_age→max_age)
    # -------------------------------------------------------------
    df_ga_f["pop_age_range"] = df_ga_f[selected_age_cols].sum(axis=1)
    pop_kom_derived = df_ga_f.groupby("Komuna")["pop_age_range"].sum()
    pop_kom_derived = pop_kom_derived[pop_kom_derived > 0]

    if pop_kom_derived.empty:
        return pd.DataFrame()

    # -------------------------------------------------------------
    # Helper për të vendosur një dimension në tabelën finale
    # -------------------------------------------------------------
    def append_block(dim, series):
        s = series[series > 0]
        if len(s) < 2:
            return
        total = s.sum()
        weights = s / total
        for cat in s.index:
            out.append({
                "Dimensioni": dim,
                "Kategoria": cat,
                "Populacioni": float(s[cat]),
                "Pesha": float(weights[cat])
            })

    # -------------------------------------------------------------
    # 3) Dimensioni: Komuna (always derived population)
    # -------------------------------------------------------------
    append_block("Komuna", pop_kom_derived)

    # -------------------------------------------------------------
    # 4) Dimensioni: Regjion (sum of derived komuna)
    # -------------------------------------------------------------
    reg_totals = {}

    for kom, popv in pop_kom_derived.items():
        reg = region_map.get(kom)
        if reg:
            reg_totals.setdefault(reg, 0)
            reg_totals[reg] += popv

    append_block("Regjion", pd.Series(reg_totals))

    # -------------------------------------------------------------
    # 5) ETHNIC SHARES (df_eth in long format) - FIXED
    # -------------------------------------------------------------

    # Step 2: Get the SAME coefficients used in sample allocation
    coef_by_komuna_for_weights = compute_gender_age_coefficients(
        df_ga,
        age_cols=age_cols,
        selected_genders=gender_selected,
        min_age=min_age,
        max_age=max_age
    )

    # Step 3: Apply coefficient to each ethnicity's base population
    df_eth_f_adj = df_eth_f.copy()
    df_eth_f_adj["coef"] = df_eth_f_adj["Komuna"].map(coef_by_komuna_for_weights).fillna(0.0)
    df_eth_f_adj["Pop_adj"] = df_eth_f_adj["Pop_base"] * df_eth_f_adj["coef"]

    # Step 4: Sum by ethnicity across all municipalities
    derived_eth_totals = df_eth_f_adj.groupby("Etnia")["Pop_adj"].sum().to_dict()

    # Fill in any missing ethnicities with 0
    eth_categories = sorted(df_eth_f["Etnia"].unique())
    for eth in eth_categories:
        if eth not in derived_eth_totals:
            derived_eth_totals[eth] = 0

    append_block("Etnia", pd.Series(derived_eth_totals))

    # -------------------------------------------------------------
    # 6) SETTLEMENT SHARES (df_eth long format)
    # -------------------------------------------------------------
    settlement_shares = {}

    for kom in df_eth_f["Komuna"].unique():
        dfk = df_eth_f[df_eth_f["Komuna"] == kom]
        vb_group = dfk.groupby("Vendbanimi")["Pop_base"].sum()

        if vb_group.sum() == 0:
            continue

        settlement_shares[kom] = (vb_group / vb_group.sum()).to_dict()

    derived_vb_totals = {vb: 0 for vb in df_eth_f["Vendbanimi"].unique()}

    for kom, popk in pop_kom_derived.items():
        if kom not in settlement_shares:
            continue
        for vb, share in settlement_shares[kom].items():
            derived_vb_totals[vb] += popk * share

    append_block("Vendbanimi", pd.Series(derived_vb_totals))

    # -------------------------------------------------------------
    # 7) GENDER SHARES for age range
    # -------------------------------------------------------------
    # compute gender-specific totals per commune
    gender_shares = {}

    for kom in df_ga_f["Komuna"].unique():
        dfk = df_ga[df_ga["Komuna"] == kom].copy()
        dfk["pop"] = dfk[selected_age_cols].sum(axis=1)
        total = dfk["pop"].sum()
        if total == 0:
            continue
        gtot = dfk.groupby("Gjinia")["pop"].sum()
        gender_shares[kom] = (gtot / total).to_dict()

    gender_categories = df_ga_f["Gjinia"].unique()
    derived_gender_totals = {g: 0 for g in gender_categories}

    for kom, popk in pop_kom_derived.items():
        if kom not in gender_shares:
            continue
        for g, share in gender_shares[kom].items():
            derived_gender_totals[g] += popk * share

    append_block("Gjinia", pd.Series(derived_gender_totals))

    # -------------------------------------------------------------
    # 8) AGE GROUPS (dynamic bins)
    # -------------------------------------------------------------
    merged_bins, labels = create_dynamic_age_groups(min_age, max_age, data_collection_method)
    derived_age_totals = {label: 0 for label in labels}

    for kom, popk in pop_kom_derived.items():

        dfk = df_ga[df_ga["Komuna"] == kom].copy()
        dfk["pop"] = dfk[selected_age_cols].sum(axis=1)
        total = dfk["pop"].sum()
        if total == 0:
            continue

        group_counts = {label: 0 for label in labels}

        for _, row in dfk.iterrows():
            for col in selected_age_cols:
                age = int(col)
                v = row[col]
                for (lo, hi), label in zip(merged_bins, labels):
                    if lo <= age <= hi:
                        group_counts[label] += v
                        break

        # now derive
        for label in labels:
            share = group_counts[label] / total
            derived_age_totals[label] += popk * share

    append_block("Grupmosha", pd.Series(derived_age_totals))

    # -------------------------------------------------------------
    # 9) APPLY CONTROLLED ROUNDING TO Populacioni PER DIMENSION
    # -------------------------------------------------------------
    if not out:
        return pd.DataFrame()

    df_final = pd.DataFrame(out)

    rounded_blocks = []

    for dim in df_final["Dimensioni"].unique():

        d = df_final[df_final["Dimensioni"] == dim].copy()

        pop_vec = d["Populacioni"].values
        total_pop = pop_vec.sum()

        # Apply your own controlled rounding function
        rounded = controlled_rounding(pop_vec, total_pop)

        d["Populacioni"] = rounded

        # Recompute weights using rounded populations
        d["Pesha"] = d["Populacioni"] / d["Populacioni"].sum()

        rounded_blocks.append(d)

    df_final = pd.concat(rounded_blocks, ignore_index=True)

    return df_final


def add_codes_to_coef_df(coef_df, data_collection_method):
    """
    Shton kolonën 'Kodi' në coef_df.
    - Të gjitha dimensionet marrin kodet fikse (si më parë)
    - Vetëm Grupmosha merr kodim dinamik sipas renditjes së saj reale (pas filtrimit)
    """

    # ======================
    # 1. Kodet fikse
    # ======================

    komuna_codes = {
        "Prishtinë":1, "Deçan":2, "Dragash":3, "Ferizaj":4, "Fushë Kosovë":5, 
        "Gjakovë":6, "Gjilan":7, "Gllogoc":8, "Graçanicë":9, "Han i Elezit":10,
        "Istog":11, "Junik":12, "Kaçanik":13, "Kamenicë":14, "Klinë":15,
        "Kllokot":16, "Leposaviq":17, "Leposavic":17, "Lipjan":18, 
        "Malishevë":19, "Mamushë":20, "Mitrovicë":21, "Mitrovica Veriore":22,
        "Novobërdë":23, "Obiliq":24, "Partesh":25, "Pejë":26, "Podujevë":27,
        "Prizren":28, "Rahovec":29, "Ranillug":30, "Skënderaj":31,
        "Suharekë":32, "Shtërpcë":33, "Shtime":34, "Viti":35, "Vushtrri":36,
        "Zubin Potok":37, "Zvecan":38
    }

    region_codes = {
        "Prishtinë":1, "Mitrovicë":2, "Pejë":3, "Prizren":4,
        "Ferizaj":5, "Gjilan":6, "Gjakovë":7
    }

    vb_codes = {"Urban":1, "Rural":2}
    gender_codes = {"Femra":1, "Femer":1, "Mashkull":2, "Meshkuj":2}
    eth_codes = {"Shqiptar":1, "Serb":2, "Tjerë":3, "Tjeter":3}

    # ==========================
    # 2. Fillimisht vendos kodet fikse
    # ==========================

    def get_fixed_code(row):
        dim = row["Dimensioni"]
        cat = row["Kategoria"]

        if dim == "Komuna": return komuna_codes.get(cat)
        if dim == "Regjion": return region_codes.get(cat)
        if dim == "Vendbanimi": return vb_codes.get(cat)
        if dim == "Gjinia": return gender_codes.get(cat)
        if dim == "Etnia": return eth_codes.get(cat)

        # Grupmosha KALO HETU – do mbushet më vonë dinamikisht
        return None

    coef_df["Kodi"] = coef_df.apply(get_fixed_code, axis=1)

    # ==========================
    # 3. KODIMI DINAMIK I GRUPMOSHËS
    # ==========================

    df_age = coef_df[coef_df["Dimensioni"] == "Grupmosha"].copy()

    if not df_age.empty:

        # (a) Parsimi i vlerave të moshës në numra
        parsed = []
        for g in df_age["Kategoria"]:
            g = str(g)

            if "-" in g:
                lo, hi = g.split("-")
                lo, hi = int(lo), int(hi)
            elif g.endswith("+"):
                lo = int(g.replace("+", ""))
                hi = 999
            else:
                continue

            parsed.append((g, lo, hi))

        # (b) Rendit grupmoshat sipas moshës
        parsed_sorted = sorted(parsed, key=lambda x: x[1])

        # (c) Gjenero kodet 1,2,3,... automatikisht
        dynamic_age_codes = {grp: i+1 for i, (grp, _, _) in enumerate(parsed_sorted)}

        # (d) Mbishkruaj kolonën Kodi *vetëm për Grupmoshën*
        coef_df.loc[coef_df["Dimensioni"] == "Grupmosha", "Kodi"] = \
            coef_df.loc[coef_df["Dimensioni"] == "Grupmosha", "Kategoria"].map(dynamic_age_codes)

    return coef_df

def create_dynamic_age_groups(age_min, age_max, data_collection_method):
    """
    Creates dynamic age bins that start at `age_min` instead of fixed 18.
    Handles merging of too-small bins automatically.
    """

    # -----------------------------------------
    # 1. Determine default base boundaries
    # -----------------------------------------
    if data_collection_method == "CAWI":
        base = [(18,24), (25,34), (35,44), (45,54), (55,200)]
    else:
        base = [(18,24), (25,34), (35,44), (45,54), (55,64), (65,200)]

    # -----------------------------------------
    # 2. Adjust first bin start to min_age
    # -----------------------------------------
    # Example: min_age=15 → change (18,24) to (15,24)

    base_adj = []
    first_lo, first_hi = base[0]
    base_adj.append((min(age_min, first_lo), first_hi))

    # Continue with remaining bins
    for lo, hi in base[1:]:
        base_adj.append((lo, hi))

    # -----------------------------------------
    # 3. Clip upper bound
    # -----------------------------------------
    if age_max is None:
        age_max = 200

    clipped = []
    for lo, hi in base_adj:
        new_lo = max(lo, age_min)
        new_hi = min(hi, age_max)
        if new_lo <= new_hi:
            clipped.append((new_lo, new_hi))

    # -----------------------------------------
    # 4. Merge bins that are smaller than 5 years
    # -----------------------------------------
    merged = []
    for lo, hi in clipped:
        if merged:
            plo, phi = merged[-1]
            if (hi - lo + 1) < 5:
                merged[-1] = (plo, hi)
            else:
                merged.append((lo, hi))
        else:
            merged.append((lo, hi))

    # Final check for last bin
    if len(merged) >= 2:
        lo, hi = merged[-1]
        if (hi - lo + 1) < 5:
            plo, phi = merged[-2]
            merged[-2] = (plo, hi)
            merged = merged[:-1]

    # -----------------------------------------
    # 5. Labels
    # -----------------------------------------
    labels = []
    for lo, hi in merged:
        if hi >= 200:
            labels.append(f"{lo}+")
        else:
            labels.append(f"{lo}-{hi}")

    return merged, labels

def generate_recode_age_dynamic(merged_bins, labels):
    """
    Gjeneron sintaksën SPSS në formatin Visual Binning.
    """

    out = "* Visual Binning.\n*Mosha.\n"
    out += "RECODE D2 (MISSING=COPY) "

    for idx, (lo, hi) in enumerate(merged_bins, start=1):
        if idx < len(merged_bins):
            out += f"(LO THRU {hi}={idx}) "
        else:
            out += f"(LO THRU HI={idx}) "

    out += "(ELSE=SYSMIS) INTO Grupmosha.\n"
    out += "VARIABLE LABELS Grupmosha 'Mosha (Binned)'.\n"
    out += "FORMATS Grupmosha (F5.0).\n"

    out += "VALUE LABELS Grupmosha "
    for idx, label in enumerate(labels, start=1):
        if "-" in label:
            a, b = label.split("-")
            label_clean = f"{a.strip()} - {b.strip()}"
        else:
            label_clean = label
        out += f"{idx} '{label_clean}' "
    out += ".\n"

    out += "VARIABLE LEVEL Grupmosha (ORDINAL).\n"
    out += "EXECUTE.\n\n"

    return out

def generate_spss_syntax(coef_df, recode_d3_text, data_collection_method):
    
    """
    Gjeneron tekstin komplet të SPSS syntax duke përfshirë:
    - RECODE D3 (siç e jep përdoruesi)
    - RECODE për Grupmosha
    - SPSSINC RAKE me të gjitha dimensionet
    """

    # --------------------------------------------
    # 1. HEADER
    # --------------------------------------------
    out = "* Encoding: UTF-8.\n\n"

    # --------------------------------------------
    # 2. Shto RECODE D3 (Regjioni) siç është dhënë
    # --------------------------------------------
    out += recode_d3_text.strip() + "\n\n"

    # --------------------------------------------
    # 3. RECODE për Grupmoshat (D2) — standarde ose dinamike
    # --------------------------------------------
    merged_bins, labels = create_dynamic_age_groups(min_age, max_age, data_collection_method)
    out += generate_recode_age_dynamic(merged_bins, labels)


    # --------------------------------------------
    # 4. SPSSINC RAKE
    # --------------------------------------------
    out += "SPSSINC RAKE\n"

    # Dimension ordering
    dim_order = list(coef_df["Dimensioni"].unique())


    dim_index = 1

    for dim in dim_order:
        df_dim = coef_df[coef_df["Dimensioni"] == dim]

        if df_dim.empty:
            continue

        out += f"DIM{dim_index}={dim} "

        for _, row in df_dim.iterrows():
            code = int(row["Kodi"])
            coef = float(row["Pesha"])
            out += f"{code} {coef}\n"

        dim_index += 1

    out += "FINALWEIGHT=peshat.\n"

    return out

def compute_natural_allocation_from_weights(coef_df, variable, value, n_total):
    """
    Returns the natural (population-proportion) interview count 
    for a given oversample variable/value.
    """
    df_dim = coef_df[coef_df["Dimensioni"] == variable]

    if df_dim.empty:
        return None  # This variable not in weight table

    row = df_dim[df_dim["Kategoria"] == value]

    if row.empty:
        return None  # Value not found

    weight = float(row["Pesha"].values[0])
    return round(weight * n_total)

def translate(term: str) -> str:
    # Normalize
    term = str(term).strip()

    # 1. Direct dictionary lookup
    if term in TRANSLATIONS:
        return TRANSLATIONS[term]

    # 2. Fallback: return original
    return term

def generate_map_url(lat, lon, zoom):
    params = urlencode({"lat": lat, "lon": lon, "zoom": zoom})
    return f"https://your-streamlit-app-url.com/?{params}"

# Load data
try:
    df_eth = load_ethnicity_settlement_data("excel-files/ASK-2024-Komuna-Etnia-Vendbanimi.xlsx")
    df_ga, age_cols = load_gender_age_data("excel-files/ASK-2024-Komuna-Gjinia-Mosha.xlsx")
except Exception as e:
    st.error(f"Gabim gjatë leximit të fajllave: {e}")
    st.stop()

region_map = get_region_mapping()

try:
    df_psu = load_psu_data("excel-files/ASK-2024-Komuna-Vendbanim-Fshat+Qytet.xlsx")
except Exception as e:
    st.error(f"Gabim gjatë leximit të fajllit të PSU-ve: {e}")
    st.stop()

# =========================
# UI: SIDEBAR
# =========================

st.title("Dizajnimi i Mostrës Nacionale")

st.sidebar.header("Parametrat kryesorë")

# Total sample size
n_total = st.sidebar.number_input(
    "Numri total i mostrës (N)",
    min_value=1,
    value=1065,
    step=100
)

# Primary stratification
primary_level = st.sidebar.selectbox(
    "Ndarja kryesore",
    options=["Komunë", "Regjion"],
    index=0
)

# Sub-stratification (can choose Vendbanim, Etnia, or both)
sub_options = st.sidebar.multiselect(
    "Nën-ndarja (mund të zgjedhësh një ose të dyja)",
    options=["Vendbanim", "Etnia"],
    default=["Vendbanim", "Etnia"]
)

st.sidebar.markdown("---")
st.sidebar.subheader("Mbledhja e të dhënave")

data_collection_method = st.sidebar.selectbox(
    "Metoda e mbledhjes së të dhënave",
    options=["CAPI", "CATI", "CAWI"],
    index=0
)
reserve_percentage = None
reserve_ratio = None

if data_collection_method=="CAPI":
    interviews_per_psu = st.sidebar.slider(
        "Numri i intervistave për PSU",
        min_value=6,
        max_value=12,
        value=8,
        step=1
    )
    survey_label = "household"
    methodology_label = "face-to-face Computer-Assisted Personal Interviewing (CAPI)"

elif data_collection_method=="CATI":
    st.sidebar.markdown("---")

    st.sidebar.subheader("Rezervat për kontakte")

    reserve_mode = st.sidebar.radio(
        "Metoda për llogaritjen e rezervave:",
        ["Përqindje (%)", "Proporcion"],
        index=0
    )

    if reserve_mode == "Përqindje (%)":
        reserve_percentage = st.sidebar.number_input(
            "Shkruaj përqindjen e rezervave (%)",
            min_value=1,
            max_value=500,
            value=20,
            step=10
        )

    elif reserve_mode == "Proporcion":
        reserve_ratio = st.sidebar.number_input(
            "Vendos numrin për proporcion (p.sh. 2 për 2:1)",
            min_value=1,
            max_value=10,
            value=2,
            step=1
        ) 

    survey_label = "individual"
    methodology_label = "Computer-Assisted Telephone Interviewing (CATI)"

elif data_collection_method=="CAWI":
    survey_label = "individual"
    methodology_label = "Computer-Assisted Web Interviewing (CAWI)"

st.sidebar.markdown("---")

# Demographic filters
st.sidebar.subheader("Filtrat demografikë")

# -----------------------------------------
# Regjioni filter
# -----------------------------------------
regjioni_filter = st.sidebar.multiselect(
    "Regjionet që përfshihen",
    options=["Prishtinë", "Mitrovicë", "Gjilan", "Gjakovë",
             "Ferizaj", "Prizren", "Pejë"],
    default=["Prishtinë", "Mitrovicë", "Gjilan", "Gjakovë",
             "Ferizaj", "Prizren", "Pejë"]
)

# -----------------------------------------
# INIT session state
# -----------------------------------------
if "komuna_filter" not in st.session_state:
    st.session_state["komuna_filter"] = sorted(df_eth["Komuna"].unique())

if "last_regjioni_filter" not in st.session_state:
    st.session_state["last_regjioni_filter"] = None

# -----------------------------------------
# APPLY REGION → KOMUNA FILTERING
# -----------------------------------------
if regjioni_filter:
    allowed_komuna = sorted(
        k for k, r in region_map.items()
        if r in regjioni_filter
    )
else:
    allowed_komuna = sorted(df_eth["Komuna"].unique())

# -----------------------------------------
# RESET vs FILTER logic (THIS IS THE FIX)
# -----------------------------------------
if st.session_state["last_regjioni_filter"] != regjioni_filter:
    # Regjioni changed → RESET Komuna
    st.session_state["komuna_filter"] = allowed_komuna.copy()
    st.session_state["last_regjioni_filter"] = regjioni_filter.copy()

else:
    # Same Regjioni → only remove invalid Komuna
    st.session_state["komuna_filter"] = [
        k for k in st.session_state["komuna_filter"]
        if k in allowed_komuna
    ]

# -----------------------------------------
# Komuna widget (NO default!)
# -----------------------------------------
komuna_filter = st.sidebar.multiselect(
    "Komunat që përfshihen",
    options=allowed_komuna,
    key="komuna_filter"
)

gender_selected = st.sidebar.multiselect(
    "Gjinia që përfshihet",
    options=["Meshkuj", "Femra"],
    default=["Meshkuj", "Femra"]
)

min_age = st.sidebar.number_input(
    "Mosha minimale (obligative)",
    min_value=0,
    value=18,
    step=1
)

max_age = st.sidebar.text_input(
    "Mosha maksimale (opsionale — lëre bosh nëse nuk ka kufi)"
)

max_age = int(max_age) if max_age.strip() else None

# Ethnicity filter (these also act as possible sub-dimensions if Etnia selected)
eth_filter = st.sidebar.multiselect(
    "Etnitë që përfshihen",
    options=["Shqiptar", "Serb", "Tjerë"],
    default=["Shqiptar", "Serb", "Tjerë"], 
    key = "Etnia-nacionale"
)

# Settlement filter
settlement_filter = st.sidebar.multiselect(
    "Vendbanimi që përfshihet",
    options=["Urban", "Rural"],
    default=["Urban", "Rural"]

)

# Oversampling
st.sidebar.markdown("---")

oversample_enabled = st.sidebar.checkbox("Oversampling", value=False)

oversample_inputs = {}

if oversample_enabled:

    oversample_vars = st.sidebar.multiselect(
        "Zgjidh deri në 2 variabla për oversample:",
        options=["Regjion", "Komuna", "Vendbanimi", "Gjinia", "Etnia", "Mosha"],
        max_selections=2
    )

    for var in oversample_vars:
        st.sidebar.markdown(f"**{var}**")

        # ============================
        # 1) SPECIAL CASE: MOSHA
        # ============================
        if var == "Mosha":
            min_over_age = st.sidebar.number_input(
                f"Grupmosha minimale ({var})",
                min_value=0, value=18, step=1,
                key=f"min_{var}"
            )
            max_over_age = st.sidebar.number_input(
                f"Grupmosha maksimale ({var})",
                min_value=min_over_age, value=24, step=1,
                key=f"max_{var}"
            )
            oversample_n = st.sidebar.number_input(
                f"Numri i anketave për {min_over_age}–{max_over_age}",
                min_value=1, value=50, step=10,
                key=f"n_{var}"
            )

            oversample_inputs[var] = [{
                "min_age": min_over_age,
                "max_age": max_over_age,
                "n": oversample_n
            }]
            continue

        # ============================
        # 2) Merr opsionet e vlefshme
        # ============================
        df_tmp = df_eth[
            (df_eth["Etnia"].isin(eth_filter)) &
            (df_eth["Vendbanimi"].isin(settlement_filter)) &
            (df_eth["Komuna"].isin(komuna_filter))
        ]

        if var == "Komuna":
            valid_kom = df_tmp.groupby("Komuna")["Pop_base"].sum()
            valid_kom = valid_kom[valid_kom > 0].index.tolist()
            options = sorted(valid_kom)
            allow_multiple = True

        elif var == "Etnia":
            options = sorted(eth_filter)
            allow_multiple = True

        elif var == "Regjion":
            valid_kom = df_tmp.groupby("Komuna")["Pop_base"].sum()
            valid_kom = valid_kom[valid_kom > 0].index.tolist()
            valid_reg = {region_map[k] for k in valid_kom if k in region_map}
            options = sorted(valid_reg)
            allow_multiple = False

        elif var == "Vendbanimi":
            options = sorted(settlement_filter)
            allow_multiple = False

        elif var == "Gjinia":
            options = sorted(gender_selected)
            allow_multiple = False

        else:
            options = []
            allow_multiple = False

        # ============================
        # 3) UI: multiselect vetëm për Komuna/Etnia
        # ============================
        if allow_multiple:
            selected_values = st.sidebar.multiselect(
                f"Zgjidh {var} për oversample (Mund të zgjidhni më shumë se një)",
                options=options,
                key=f"multi_{var}"
            )

            entry_list = []
            for v in selected_values:
                q = st.sidebar.number_input(
                    f"Kuota për {var} = {v}",
                    min_value=1, value=50, step=10,
                    key=f"quota_{var}_{v}"
                )
                entry_list.append({"value": v, "n": q})

            oversample_inputs[var] = entry_list

        else:
            val = st.sidebar.selectbox(
                f"Njësia nga {var} që do të jetë oversample",
                options=options,
                key=f"val_{var}"
            )
            q = st.sidebar.number_input(
                f"Kuota për {var} = {val}",
                min_value=1, value=50, step=10,
                key=f"quota_{var}_{val}"
            )

            oversample_inputs[var] = [{"value": val, "n": q}]


st.sidebar.markdown("---")

seed = 42

st.sidebar.markdown("Kliko më poshtë për të llogaritur shpërndarjen.")
run_button = st.sidebar.button("Gjenero shpërndarjen e mostrës")

# =========================
# MAIN LOGIC
# =========================

if run_button:

    # 1) Filter ethnicity & settlement (these are demographic filters)
    df = df_eth.copy()
    df = df[df["Etnia"].isin(eth_filter)]
    df = df[df["Vendbanimi"].isin(settlement_filter)]
    df = df[df["Komuna"].isin(komuna_filter)]

    if df.empty:
        st.error("Asnjë kombinim nuk përputhet me filtrat e zgjedhur (Etnia/Vendbanimi).")
        st.stop()

    # 2) Compute gender+age coefficients per Komuna
    coef_by_komuna = compute_gender_age_coefficients(
        df_ga,
        age_cols=age_cols,
        selected_genders=gender_selected,
        min_age=min_age,
        max_age=max_age
    )

    # Mapojmë gjininë origjinale në df_eth duke përdorur df_ga si referencë
    gender_map = df_ga.groupby("Komuna")["Gjinia"].apply(list).to_dict()

    df["Gjinia"] = df["Komuna"].map(lambda k: gender_map.get(k, ["Meshkuj","Femra"]))
                    
    # zgjerim i rreshtave për çdo gjini (ndryshe OS nuk punon)
    df = df.explode("Gjinia")

    # Attach coefficient to df (missing komuna -> coef 0)
    df["coef_gender_age"] = df["Komuna"].map(coef_by_komuna).fillna(0.0)

    # 3) Adjusted population for filters (Etnia, Vendbanim, Gjinia, Mosha)
    df["Pop_adj"] = df["Pop_base"] * df["coef_gender_age"]

    # Remove rows with zero adjusted population
    df = df[df["Pop_adj"] > 0]

    if df.empty:
        st.error("Pas aplikimit të koeficientëve (gjinia/mosha), Pop_adj është 0 për të gjitha njësitë.")
        st.stop()

    # 4) Primary stratification
    if primary_level == "Regjion":
        if not region_map:
            st.warning(
                "Ndarja sipas Regjionit kërkon të plotësohet 'region_map' në kod. "
                "Aktualisht nuk ka mapping, prandaj po vazhdohet vetëm me nivel Komune."
            )
            base_col = "Komuna"
        else:
            df["Regjion"] = df["Komuna"].map(region_map)
            df = df.dropna(subset=["Regjion"])
            base_col = "Regjion"
    else:
        base_col = "Komuna"

    # 5) Sub-stratification labels
    # Ensure consistent sorting at ethnicity
    eth_order = ["Shqiptar", "Serb", "Tjerë"]
    df["Etnia"] = pd.Categorical(df["Etnia"], categories=eth_order, ordered=True)

    # Combine ethnicity with settlement
    if "Etnia" in sub_options and "Vendbanim" in sub_options:
        df["Sub"] = df["Etnia"].astype(str) + " - " + df["Vendbanimi"].astype(str)
    elif "Etnia" in sub_options:
        df["Sub"] = df["Etnia"].astype(str)
    elif "Vendbanim" in sub_options:
        df["Sub"] = df["Vendbanimi"].astype(str)
    else:
        df["Sub"] = "Total"
    
    if not sub_options:
        df["Sub"] = df["Vendbanimi"].astype(str)

    grouped = (
        df.groupby([base_col, "Sub"], as_index=False)["Pop_adj"]
        .sum()
        .rename(columns={"Pop_adj": "Pop_stratum"})
    )

    grouped = grouped.reset_index(drop=True)

    precomputed_masks = {}
    for var, entry_list in oversample_inputs.items():
        for entry in entry_list:
            precomputed_masks[var] = mask_for_oversample(grouped, var, entry)

    # Sort columns
    sub_order = []
    for eth in eth_order:
        for vb in ["Urban", "Rural"]:
            sub_order.append(f"{eth} - {vb}")
    sub_order += eth_order + ["Urban", "Rural", "Total"]  # për raste të tjera

    # Filtwe columns
    existing_subs = sorted(grouped["Sub"].unique(), key=lambda x: sub_order.index(x) if x in sub_order else 999)
    grouped["Sub"] = pd.Categorical(grouped["Sub"], categories=existing_subs, ordered=True)


    total_pop = grouped["Pop_stratum"].sum()
    if total_pop <= 0:
        st.error("Popullsia totale pas filtrave është 0. Nuk mund të alokohet mostra.")
        st.stop()

    # =====================
    #  Margin of Error (95%)
    # =====================

    z = 1.96           # 95% confidence
    p = 0.5            # worst-case scenario
    n = n_total
    N = total_pop      # from grouped Pop_stratum sum

    if N > n:
        fpc = ((N - n) / (N - 1)) ** 0.5
    else:
        fpc = 1.0

    moe = z * ((p * (1 - p)) / n) ** 0.5 * fpc
    moe_percent = moe * 100


    all_os = []

    for var, entries in oversample_inputs.items():
        for entry in entries:

            mask = mask_for_oversample(grouped, var, entry)

            # MOSHA ka strukturë ndryshe → nuk ka "value"
            if var == "Mosha":
                all_os.append({
                    "var": var,
                    "value": f"{entry['min_age']}-{entry['max_age']}",
                    "n": entry["n"],
                    "mask": mask
                })
            else:
                # Komuna, Etnia, Regjioni, Vendbanimi, Gjinia
                all_os.append({
                    "var": var,
                    "value": entry["value"],
                    "n": entry["n"],
                    "mask": mask
                })

    # ================================
    # 7a) Oversampling
    # ================================
    # Save natural allocation BEFORE oversampling
    grouped["n_alloc"] = 0

    oversample_items = list(oversample_inputs.items())

    # Helper për alokim proporcional në një maskë
    def alloc_to_mask(mask: pd.Series, quota: int) -> pd.Series:
        out = pd.Series(0.0, index=grouped.index)
        quota = int(quota)
        if quota <= 0:
            return out

        df_sub = grouped[mask].copy()
        if df_sub.empty:
            return out

        weights = df_sub["Pop_stratum"] / df_sub["Pop_stratum"].sum()
        floats = weights * quota
        ints = controlled_rounding(floats.to_numpy(), quota, seed)
        out.loc[df_sub.index] = ints
        return out

    # 0) Nuk ka oversample fare → alokim i thjeshtë proporcional
    if not oversample_items:
        total_pop = grouped["Pop_stratum"].sum()
        weights = grouped["Pop_stratum"] / total_pop
        floats = weights * n_total
        grouped["n_alloc"] = controlled_rounding(floats.to_numpy(), n_total, seed)

    # 1) Vetëm një oversample
    elif len(oversample_items) == 1:
        varA, entry_list = oversample_items[0]

        # fillo me zero
        grouped["n_alloc"] = 0

        # për çdo kuotë të atij variabli
        used_total = 0
        for entry in entry_list:
            nA = int(entry["n"])
            maskA = mask_for_oversample(grouped, varA, entry)

            alloc_A = alloc_to_mask(maskA, nA)
            grouped["n_alloc"] += alloc_A
            used_total += int(alloc_A.sum())

        # pjesa e mbetur shkon tek stratat tjera
        remaining = n_total - used_total
        if remaining < 0:
            remaining = 0

        mask_rest = (grouped["n_alloc"] == 0)
        alloc_rest = alloc_to_mask(mask_rest, remaining)

        grouped["n_alloc"] += alloc_rest

    # 2) Dy variabla oversample (me shumë vlera për njërin variabël)
    else:
        # ndërto listën e plotë që i ke më lart
        # all_os = [ {var, value, n, mask}, ... ]

        # Renditi sipas kuotës zbritëse
        all_os_sorted = sorted(all_os, key=lambda x: x["n"], reverse=True)

        # OS_B = variabli me kuotën më të lartë (p.sh. Rural=800)
        osB = all_os_sorted[0]

        # OS_A = të gjithë tjerët (p.sh. Peja=300, Prishtina=500, Gjakova=200)
        osA_list = all_os_sorted[1:]

        # 1) shpërndaj OS_B
        alloc_B = alloc_to_mask(osB["mask"], osB["n"])

        grouped["n_alloc"] = alloc_B

        # 2) pastaj secilin OS_A një nga një
        for osA in osA_list:

            # llogarit overlapp
            overlap_mask = osA["mask"] & osB["mask"]
            overlap_from_B = int(alloc_B[overlap_mask].sum())

            remaining_A = max(osA["n"] - overlap_from_B, 0)

            alloc_A = alloc_to_mask(osA["mask"] & ~osB["mask"], remaining_A)

            grouped["n_alloc"] += alloc_A

        # 3) pjesa e mbetur shkon jashtë OS-ve
        used = int(grouped["n_alloc"].sum())
        remaining = max(n_total - used, 0)

        mask_rest = ~(sum(os["mask"] for os in all_os) > 0)
        alloc_rest = alloc_to_mask(mask_rest, remaining)

        grouped["n_alloc"] += alloc_rest

    # ================================
    # 8) Heq kolonat që s'duhen para pivot-it
    # ================================
    drop_cols = []
    if "Gjinia" in grouped.columns:
        drop_cols.append("Gjinia")
    if "AgeSeg" in grouped.columns:
        drop_cols.append("AgeSeg")

    if drop_cols:
        grouped = (
            grouped.groupby([base_col, "Sub"])[["Pop_stratum", "n_alloc"]]
                .sum()
                .reset_index()
        )

    # 9) Prepare pivot table: rows = primary, columns = sub-dimensions
    pivot = grouped.pivot(
        index=base_col,
        columns="Sub",
        values="n_alloc"
    ).fillna(0).astype(int)

    original_strata_cols = list(pivot.columns)

    # Add total per primary
    pivot["Total"] = pivot[original_strata_cols].sum(axis=1)
    
    # ==========================================================
    #  OVERSAMPLE GENDER/MOSHA pas pivot (në nivel KOMUNE)
    # ==========================================================

     # Marrim TOTAL-in e komunës nga pivot
    pivot_totals = pivot["Total"].copy()

        # -------------------------
        # 1) GJINIA
        # -------------------------
    if "Gjinia" in oversample_inputs:

        os_gender = oversample_inputs["Gjinia"][0]["value"]
        os_n = int(oversample_inputs["Gjinia"][0]["n"])

        if os_n > n_total:
            st.warning(
            f"Vërejtje: Kuota e alokuar ({os_n}) për oversample tek ({os_gender}) është më e madhe se N = ({n_total}). "
            "Shëno një kuotë tjetër për oversample."
        )

        # Popullsia sipas gjinisë per komunë
        pop_by_gender = (
            df_ga.groupby(["Komuna", "Gjinia"])[age_cols]
            .sum()
            .sum(axis=1)
            .unstack(fill_value=0)
            .reindex(pivot.index)
            .fillna(0)
            )

        pop_os = pop_by_gender[os_gender]
        weight_os = pop_os / pop_os.sum()

        # Alokimi për OS
        os_alloc = (weight_os * os_n).round().astype(int)

        # Alokimi për pjesën tjetër
        leftover = pivot_totals - os_alloc

        if os_gender == "Femra":
            pivot["Femra"] = controlled_rounding(os_alloc, os_n)
            pivot["Meshkuj"] = pivot_totals - pivot["Femra"]
        else:
            pivot["Meshkuj"] = controlled_rounding(os_alloc, os_n)
            pivot["Femra"] = pivot_totals - pivot["Meshkuj"]

        # -------------------------
        # 2) MOSHA
        # -------------------------
    if "Mosha" in oversample_inputs:

        params_age = oversample_inputs["Mosha"][0]
        os_min = params_age["min_age"]
        os_max = params_age["max_age"]
        os_n = int(params_age["n"])

        if os_n > n_total:
            st.warning(
            f"Vërejtje: Kuota e alokuar ({os_n}) për oversample tek ({os_min}-{os_max}) është më e madhe se N = ({n_total}). "
            "Shëno një kuotë tjetër për oversample."
        )

        # Lista e moshave që ekzistojnë në dataset
        age_cols_sorted = sorted(age_cols, key=lambda x: int(str(x)))

        # Grupi OS (18–30 p.sh.)
        range_os = [c for c in age_cols_sorted if os_min <= int(c) <= os_max]

        # Grupi jashtë OS (31+ p.sh.)
        range_non = [c for c in age_cols_sorted if int(c) > os_max]

        # Popullsia sipas moshës per komunë
        pop_by_age = (
            df_ga.groupby("Komuna")[age_cols_sorted]
            .sum()
            .reindex(pivot.index)
            .fillna(0)
            )

        pop_os = pop_by_age[range_os].sum(axis=1)
        pop_non = pop_by_age[range_non].sum(axis=1)

        # Pesha për OS
        weight_os_age = pop_os / pop_os.sum()

        # 1) SHPËRNDA OS për moshë
        os_alloc_age = (weight_os_age * os_n).round().astype(int)

        # 2) Alokimi final për moshë
        age_label_os = f"{os_min}–{os_max}"
        age_label_non = f"{os_max+1}+"

        pivot[age_label_os] = controlled_rounding(os_alloc_age, os_n)
        pivot[age_label_non] = pivot_totals - pivot[age_label_os]

    pivot_old = pivot.copy()
    
    ###########################################
    # FIX MAJORITY ETHNICITY CALCULATION HERE
    ###########################################
    if "Etnia" in sub_options:
        eth_majority = {}

        for kom in pivot.index:
            totals = {}
            
            # Find all ethnicity groups dynamically
            for eth in ["Shqiptar", "Serb", "Tjerë"]:
                cols = [c for c in pivot.columns if c.startswith(eth)]
                total = sum(pivot.at[kom, c] for c in cols)
                totals[eth] = total

            # If all totals are 0 → no majority
            if all(v == 0 for v in totals.values()):
                eth_majority[kom] = None
            else:
                eth_majority[kom] = max(totals, key=totals.get)

        majority = eth_majority

    else:
        majority = {}

        ###########################################
        # FIX MAJORITY ETHNICITY CALCULATION HERE
        ###########################################

    pivot = fix_minimum_allocations(
            pivot=pivot,
            df_eth= df_eth,
            region_map=region_map,
            strata_col = original_strata_cols,
            majority=majority,
            selected_ethnicity=eth_filter,
            min_total=3,   # minimum anketa per komunë
            min_eth=3      # minimum per vendbanim (Urban/Rural)
            )
    
    pivot = pivot[pivot["Total"] != 0]

    # ==========================================================
    # RECALCULATE NON-OVERSAMPLED CATEGORIES AFTER FIX ALLOCATION
    # ==========================================================

    for var, entries in oversample_inputs.items():

        # Skip variables with multiple oversampled categories (Komuna/Etnia)
        if len(entries) != 1:
            continue

        if var != "Mosha":
            os_value = entries[0]["value"]

        # GJINIA
        if var == "Gjinia":
            # Which category is NOT oversampled?
            other_gender = "Femra" if os_value == "Meshkuj" else "Meshkuj"

            # Get the oversample column (created earlier)
            if os_value in pivot.columns:
                pivot[other_gender] = pivot["Total"] - pivot[os_value]

        # MOSHA
        elif var == "Mosha":
            os_label = f"{entries[0]['min_age']}–{entries[0]['max_age']}"
            non_label = f"{entries[0]['max_age']+1}+"

            if os_label in pivot.columns:
                pivot[non_label] = pivot["Total"] - pivot[os_label]

        # KOMUNA / REGJION — no need to fix, because these do NOT create new columns
        
    if not sub_options:
        pivot = pivot.drop(columns=["Urban"])
        pivot = pivot.drop(columns=["Rural"])
    
    # Safety: ensure global total matches n_total
    global_total = int(pivot.loc["Total", "Total"])

    # Përgatit tekstin për grupmoshën
    if max_age is None:
        age_text = f"{min_age}+"
    else:
        age_text = f"{min_age}–{max_age}"

    # Përgatit tekstin për gjininë
    if len(gender_selected) == 1:
        gender_text = f"{gender_selected[0]}"
    else:
        gender_text = "Femra, Meshkuj"

    if len(settlement_filter) == 1:
        settlement_text = f"{settlement_filter[0]}"
    else:
        settlement_text = "Urban, Rural"

    if len(eth_filter) == 1 or len(eth_filter) == 2:
        ethnicity_text = ", ".join(eth_filter)
    else:
        ethnicity_text = "Shqiptar, Serb, Tjerë"

    if oversample_enabled:
        parts = []
        for var, entries in oversample_inputs.items():
            values = []
            for e in entries:
                if var == "Mosha":
                    values.append(f"{e['min_age']}–{e['max_age']}")
                else:
                    values.append(str(e["value"]))

            # Join multiple values
            values_text = ", ".join(values)

            parts.append(f"{var} → ({values_text})")

        oversampling_text = "; ".join(parts)
    else:
        oversampling_text = "Joaktiv"


    # ============================
    # 3 CARDS UI
    # ============================
    def load_svg(path):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()

    icon_sample = load_svg("images/sample-people.svg")
    icon_strata = load_svg("images/strata.svg")
    icon_demo = load_svg("images/demographics.svg")

    col1, col2, col3 = st.columns(3)

    with col1:
        with st.container():
            st.markdown(f"""
            <div class='card'>
                <div class='card-title'>
                    {icon_sample} Mostra
                </div>
                <div class='card-value'>Totali i mostrës: <b>{n_total}</b></div>
                <div class='card-value'>Marzha e gabimit: <b>± {moe_percent:.2f}%</b></div>
                <div class='card-value'>Intervali i besimit: <b>95%</b></div>
            </div>
            """, unsafe_allow_html=True)

    with col2:
        with st.container():
            st.markdown(f"""
            <div class='card'>
                <div class='card-title'>
                        {icon_strata} Dizajnimi i Mostrës
                    </div>
                <div class='card-value'>Ndarja kryesore: <b>{primary_level}</b></div>
                <div class='card-value'>Nën-ndarja: <b>{", ".join(sub_options)}</b></div>
                <div class='card-value'>Oversampling: <b>{oversampling_text}</b></div>
                <div class='card-value'><b></b></div>
            </div>
            """, unsafe_allow_html=True)

    with col3:
        with st.container():
            st.markdown(f"""
            <div class='card'>
                <div class='card-title'>
                        {icon_demo} Demografia
                </div>
                <div class='card-value'>Grupmosha: <b>{age_text}</b></div>
                <div class='card-value'>Gjinia: <b>{gender_text}</b></div>
                <div class='card-value'>Vendbanimi: <b>{settlement_text}</b></div>
                <div class='card-value'>Etnia: <b>{ethnicity_text}</b></div>
            </div>
            """, unsafe_allow_html=True)

    st.subheader("Tabela e ndarjes së mostrës")
    
    st.dataframe(pivot, use_container_width=True)

    if global_total != n_total:
        st.warning(
            f"Vërejtje: Totali i alokuar ({global_total}) nuk përputhet me N = ({n_total}). "
            "Kontrollo koeficientët dhe numerikën."
        )

    # 10) Show long format result (optional, më teknik)
    with st.expander("Shfaq tabelën e plotë të stratum-eve (long format)", expanded=False):
        display_cols = [base_col, "Sub", "Pop_stratum", "n_alloc"]
        st.dataframe(grouped[display_cols], use_container_width=True)


    # 📘 Pivot table (Excel)
    pivot_excel = df_to_excel_bytes(pivot, sheet_name="Mostra")
    create_download_link(
        file_bytes=pivot_excel,
        filename="mostra_e_gjeneruar.xlsx",
        label="Shkarko Mostrën"
    )

    # 📘 Strata table (Excel)
    strata_excel = df_to_excel_bytes(grouped, sheet_name="Strata")
    create_download_link2(
        file_bytes=strata_excel,
        filename="mostra_strata.xlsx",
        label="Shkarko Strata"
    )

    # 📘 Old Pivot table (Excel)
    old_pivot_excel = df_to_excel_bytes(pivot_old, sheet_name="Shpërndarja fillestare")
    create_download_link2(
        file_bytes=old_pivot_excel,
        filename="shpërndarja_fillestare.xlsx",
        label="Shkarko Shpërndarjen Fillestare"
    )

    # =====================================================
    # PSU-të për CAPI – tani funksionon si për Komunë ashtu edhe për Regjion
    # =====================================================
    if data_collection_method == "CAPI":
        st.markdown("---")
        st.subheader("PSU-të e përzgjedhura")

        # --------------------------------------------
        # a) Ndërto pivot-in që do përdoret vetëm për PSU
        #    - nëse ndarja kryesore është Komuna → përdor pivot ekzistues
        #    - nëse ndarja kryesore është Regjion → ri-aloko N në nivel komune
        # --------------------------------------------
        if primary_level == "Komunë":
            pivot_for_psu = pivot.copy()
        else:
            # për Regjion, llogarisim një shpërndarje të thjeshtë të N në nivel Komune
            # në bazë të Pop_adj (pas filtrave)
            df_mun_pop = df.groupby("Komuna")["Pop_adj"].sum()

            # siguri – hiq komuna me 0 populacion
            df_mun_pop = df_mun_pop[df_mun_pop > 0]

            if df_mun_pop.empty:
                st.warning("Nuk mund të llogariten PSU sepse nuk ka popullsi valide në nivel komune.")
                psu_table = pd.DataFrame()
            else:
                weights_mun = df_mun_pop / df_mun_pop.sum()
                floats = weights_mun * n_total
                totals = controlled_rounding(floats.to_numpy(), n_total, seed)

                pivot_for_psu = pd.DataFrame(
                    {"Total": totals},
                    index=df_mun_pop.index
                )
                pivot_for_psu.loc["Total"] = pivot_for_psu.sum()
            
        with st.spinner("Duke llogaritur PSU-të..."):
            if "pivot_for_psu" in locals():
                psu_table = compute_psu_table_for_all_municipalities(
                    pivot=pivot_for_psu,
                    df_psu=df_psu,
                    k=interviews_per_psu,
                    eth_filter=eth_filter,
                    settlement_filter=settlement_filter,
                )
            else:
                psu_table = pd.DataFrame()

        if psu_table.empty:
            st.warning("Nuk u gjeneruan PSU. Kontrollo filtrat, fajllin e PSU-ve dhe shpërndarjen e mostrës.")
        else:
            st.caption(
                f"PSU-të janë llogaritur me **{interviews_per_psu} intervista** për PSU sipas rregullit të përcaktuar."
            )
            st.dataframe(psu_table, use_container_width=True)

            psu_excel = df_to_excel_bytes(psu_table, sheet_name="PSU")
            create_download_link(
                file_bytes=psu_excel,
                filename="psu_capi_tegjitha_komunat.xlsx",
                label="Shkarko PSU-të"
            )

        # =====================================================
        # Harta – gjithmonë nëse kemi psu_table, pavarësisht ndarjes kryesore
        # =====================================================
        if not psu_table.empty:
            st.subheader("Harta e PSU-ve të përzgjedhura")

            # Remove the artificial urban row BEFORE merging with coordinates
            df_map = psu_table[["Komuna", "Fshati/Qyteti", "Intervista"]].copy()
            df_map.loc[df_map["Fshati/Qyteti"] == "Urban", "Fshati/Qyteti"] = \
                df_map.loc[df_map["Fshati/Qyteti"] == "Urban", "Komuna"]

            # Merge with PSU coordinates
            df_map = df_map.merge(
                df_psu[["Komuna", "Fshati/Qyteti", "lat", "long"]],
                on=["Komuna", "Fshati/Qyteti"],
                how="left"
            )

            layer = pdk.Layer(
                "ScatterplotLayer",
                data=df_map,
                get_position='[long, lat]',
                get_fill_color='[200, 30, 0, 160]',
                get_radius=600,
                pickable=True
            )

            view_state = pdk.ViewState(
                latitude=df_map["lat"].mean(),
                longitude=df_map["long"].mean(),
                zoom=8
            )

            deck = pdk.Deck(
                layers=[layer],
                initial_view_state=view_state,
                map_provider="carto",
                map_style="light",
                tooltip={"html": "<b>{Komuna}</b><br>{Fshati/Qyteti}</b><br>{Intervista} intervista"}
            )

            st.pydeck_chart(deck)
            deck_html = deck.to_html(as_string=True)
            html_bytes = deck_html.encode("utf-8")
            # Butoni i shkarkimit
            create_download_link(html_bytes, "psu_map.html", "Shkarko hartën (HTML)")
    
    if data_collection_method == "CATI":

        st.markdown("---")
        st.subheader("Lista e kontakteve")

        df_contacts_raw = load_citizens_database()

        df_contacts_filtered = filter_contacts(
            df_contacts_raw,
            komuna_filter=komuna_filter,
            gender_selected=gender_selected,
            min_age=min_age,
            max_age=max_age,
            settlement_filter=settlement_filter
        )

        if df_contacts_filtered.empty:
            st.warning("Nuk ka kontakte që përputhen me filtrat e zgjedhur.")
        else:
            df_contacts_final = build_contact_list(
                df_contacts=df_contacts_filtered,
                pivot=pivot,
                reserve_mode=reserve_mode,
                reserve_percentage=reserve_percentage,
                reserve_ratio=reserve_ratio
            )

            if df_contacts_final.empty:
                st.warning("Nuk u gjenerua asnjë listë kontakti.")
            else:

                display_cols = [
                    "Emri dhe mbiemri",
                    "Numri i telefonit",
                    "Gjinia",
                    "Mosha",
                    "Komuna",
                    "Vendbanimi",
                    "Etnia",
                    "Edukimi",
                    "Punësimi",
                    "Statusi martesor",
                    "Të ardhurat familjare",
                    "Të ardhurat personale"
                ]

                st.dataframe(
                    df_contacts_final[display_cols],
                    use_container_width=True
                )

                excel_bytes = df_to_excel_bytes(
                    df_contacts_final[display_cols],
                    sheet_name="Kontaktet_CATI"
                )

                create_download_link(
                    excel_bytes,
                    "lista_kontakteve_cati.xlsx",
                    "Shkarko Listën e Kontakteve"
                )

    # COMMON SECTION (always included)
    if not oversample_enabled:
        third_strata = "Gender and Age Group"
    
    else:
        if "Gjinia" in oversample_vars:
            third_strata = "Age Group"
        else: 
            third_strata = "Gender"

    coef_df = compute_population_coefficients(
    df_ga=df_ga,
    df_eth=df_eth,
    region_map=region_map,
    gender_selected=gender_selected,
    min_age=min_age,
    max_age=max_age,
    eth_filter=eth_filter,
    settlement_filter=settlement_filter,
    komuna_filter=komuna_filter,
    data_collection_method=data_collection_method
    )

    coef_df = add_codes_to_coef_df(coef_df, data_collection_method)
    # Remove dimensions with only 1 category
    filtered_dims = (
        coef_df.groupby("Dimensioni")["Kategoria"]
        .nunique()
    )

    # ============================================
    # NATURAL ALLOCATIONS for OVERSAMPLING narrative
    # ============================================
    os_additional_list = []

    for var, entries in oversample_inputs.items():

        for entry in entries:

            # Category label
            if var == "Mosha":
                cat_label = f"{entry['min_age']}–{entry['max_age']}"
            else:
                cat_label = entry["value"]

            # Compute natural allocation BEFORE removing dimensions
            natural = compute_natural_allocation_from_weights(
                coef_df, variable=var, value=cat_label, n_total=n_total
            )

            # Target (oversample quota)
            target = entry["n"]

            # Added respondents
            added = max(target - (natural or 0), 0)

            # Store for narrative
            os_additional_list.append({
                "var": var,
                "value": cat_label,
                "natural": natural,
                "target": target,
                "added": added
            })

    second_level = []
    for i in sub_options:
        i = translate(i)
        second_level.append(i)

    strata_list = [primary_level] + sub_options
    narrative_text = narrative_template_common.format(
        survey_label = survey_label,
        methodology_label = methodology_label,
        n_total=n_total,
        moe= f"{round(moe * 100, 2)}%",
        primary_level = translate(primary_level),
        second_level = " and ".join(second_level),
        third_level = translate(third_strata)
    )

    # METHOD-SPECIFIC SECTION
    if data_collection_method == "CAPI":
        narrative_text += narrative_template_capi.format(
            interviews_per_psu=interviews_per_psu
        )

    elif data_collection_method == "CATI":
        narrative_text += narrative_template_cati

    if oversample_enabled:
        os_group_list = ""
        for item in os_additional_list:
            os_group_list += (
                f"- **{translate(item['value'])}**: {item['added']} additional interviews  \n"
            )
        os_target_group = "; ".join(
            f"{translate(item['value'])}" 
            for item in os_additional_list
        )

        os_added_total = ", ".join(
            translate(str(item["added"])) 
            for item in os_additional_list
        )

        if len(os_additional_list) == 1:
            narrative_text += narrative_template_oversampling_single_active.format(
                os_target_group=os_target_group,
                os_added_total=os_added_total
            )
        else:
            narrative_text += narrative_template_oversampling_multi_active.format(
                os_group_list=os_group_list
            )

    else:
        narrative_text += narrative_template_oversampling
        narrative_text += narrative_template_oversampling_inactive

    st.markdown("---")  
    st.subheader("Përshkrimi i dizajnimit të mostrës")

    with st.expander("Shfaq përshkrimin e dizajnimit të mostrës"):
        st.markdown(narrative_text)

    
    narrative_doc = narrative_to_word(narrative_text)

    b64 = base64.b64encode(narrative_doc).decode()

    st.markdown(f"""
        <a href="data:application/octet-stream;base64,{b64}" download="pershkrimi_mostres.docx">
            <div style="
                background-color:#344b77;
                color:white;
                text-align:center;
                font-weight:500;
                font-size:16px;
                padding:10px;
                border-radius:8px;
                margin-top:10px;
                cursor:pointer;">
                Shkarko Përshkrimin (Word)
            </div>
        </a>
    """, unsafe_allow_html=True)

    dims_to_keep = filtered_dims[filtered_dims > 1].index.tolist()

    coef_df = coef_df[coef_df["Dimensioni"].isin(dims_to_keep)]

    if coef_df.empty:
        st.warning("Nuk ka asnjë dimension valid për peshim pas filtrave.")
        st.stop()

    st.markdown("---")
    st.subheader("Sintaksa për peshim në SPSS")

    with st.expander("Shfaq tabelën e plotë të peshave", expanded=False):
        st.dataframe(coef_df, use_container_width=True)

    coef_df = coef_df.dropna(subset=["Kodi"])
    spss_text = generate_spss_syntax(
    coef_df,
    recode_d3_text=RECODE_D3_TEMPLATE,
    data_collection_method=data_collection_method
)

    create_download_link(
        file_bytes=spss_text.encode("utf-8"),
        filename="syntax_peshat.sps",
        label="Shkarko Peshat për SPSS"
    )

else:
    st.info("Cakto parametrat kryesorë dhe kliko **'Gjenero shpërndarjen e mostrës'** për të dizajnuar mostrën.")
